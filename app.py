from __future__ import annotations

import hashlib
import hmac
import html
import json
import os
import re
import logging
import textwrap
import uuid
from dataclasses import dataclass
from datetime import date, datetime, timedelta
import calendar as py_calendar
from io import BytesIO, StringIO
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import pandas as pd
import requests
import streamlit as st

try:
    import altair as alt
except Exception:
    alt = None

try:
    from streamlit_calendar import calendar as streamlit_calendar
except Exception:
    streamlit_calendar = None
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

try:
    from docx import Document
except Exception:
    Document = None

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

# ============================================================
# 기본 설정
# ============================================================

APP_ROOT = Path(__file__).resolve().parent
DATA_DIR = APP_ROOT / "data"
STORE_DIR = APP_ROOT / "storage"
WHITEPAPER_DIR = DATA_DIR / "whitepaper"
UPLOAD_DIR = STORE_DIR / "uploads"
DATA_DIR.mkdir(exist_ok=True)
STORE_DIR.mkdir(exist_ok=True)
WHITEPAPER_DIR.mkdir(exist_ok=True)
UPLOAD_DIR.mkdir(exist_ok=True)

APP_TITLE = "AIVLE 학습도우미"
OPENROUTER_APP_TITLE = "AIVLE Learning Assistant"
OPENROUTER_ENDPOINT = "https://openrouter.ai/api/v1/chat/completions"
DEFAULT_OPENROUTER_MODEL = "nvidia/nemotron-3-super-120b-a12b:free"
OPENROUTER_TIMEOUT_SECONDS = 30
CURRENT_WHITEPAPER_STEM = "current_whitepaper"
WHITEPAPER_EXTENSIONS = {".docx", ".pdf", ".txt"}
BUNDLED_WHITEPAPER_PATH = DATA_DIR / "aivle_kt_learning_whitepaper_2026.docx"
DEFAULT_WHITEPAPER_PATH = WHITEPAPER_DIR / f"{CURRENT_WHITEPAPER_STEM}.docx"
WHITEPAPER_META_PATH = STORE_DIR / "whitepaper_meta.json"
CONVERSATION_PATH = STORE_DIR / "conversations.json"
RESULT_PATH = STORE_DIR / "study_results.json"
WRONG_NOTE_PATH = STORE_DIR / "wrong_notes.json"
CALENDAR_PATH = STORE_DIR / "calendar_events.json"
CHECKLIST_PATH = STORE_DIR / "daily_checklist.json"
CAREER_REPORT_PATH = STORE_DIR / "career_reports.json"
logger = logging.getLogger(__name__)

try:
    if not any(WHITEPAPER_DIR.glob(f"{CURRENT_WHITEPAPER_STEM}.*")) and BUNDLED_WHITEPAPER_PATH.exists():
        DEFAULT_WHITEPAPER_PATH.write_bytes(BUNDLED_WHITEPAPER_PATH.read_bytes())
except Exception:
    pass

MENU_OPTIONS = [
    "대시보드",
    "학습 질의",
    "예습·진단",
    "복습·분석",
    "취업 준비",
    "일정·커리큘럼",
    "내 학습 현황",
]

TOPIC_ALIASES = {
    "분석형 AI": ["분석형 AI", "데이터 분석", "머신러닝", "딥러닝", "모델", "예측", "분류", "회귀"],
    "생성형 AI": ["생성형 AI", "LLM", "프롬프트", "RAG", "LangChain", "이미지 모델", "언어 모델"],
    "Cloud": ["Cloud", "클라우드", "IT Infra", "Cloud Infra", "Cloud Native", "서버", "인프라"],
    "서비스 개발/제안": ["서비스 개발", "서비스 제안", "SW", "Web", "제안서", "기획", "프로젝트 관리"],
    "프로젝트 수행": ["미니프로젝트", "빅프로젝트", "MVP", "주제 선정", "발표", "현직자 코칭", "포트폴리오"],
    "출결·운영": ["출결", "지각", "조퇴", "외출", "공가", "수료", "에이블에듀", "체크인", "체크아웃"],
}

OFFICIAL_NOTICE_KEYWORDS = ["출결", "공가", "수료", "지원 자격", "자비부담금", "평가", "합격", "불합격", "최신", "공식", "일정"]

DEFAULT_CURRICULUM = [
    {"week": "1주", "period": "3.30~4.3", "topic": "입교식 + 분석형 AI", "kind": "수업", "note": "입교 오리엔테이션, 분석형 AI 시작"},
    {"week": "2주", "period": "4.6~4.10", "topic": "분석형 AI", "kind": "수업", "note": "데이터 분석·머신러닝 기초"},
    {"week": "3주", "period": "4.13~4.17", "topic": "분석형 AI + 분석형 AI 미니프로젝트", "kind": "프로젝트", "note": "분석형 AI 실습 산출물"},
    {"week": "4주", "period": "4.20~4.24", "topic": "분석형 AI 미니프로젝트 + 생성형 AI", "kind": "프로젝트", "note": "생성형 AI 전환"},
    {"week": "5주", "period": "4.27~5.1", "topic": "생성형 AI 미니프로젝트 + 생성형 AI", "kind": "프로젝트", "note": "5.1 휴강"},
    {"week": "6주", "period": "5.4~5.8", "topic": "생성형 AI", "kind": "수업", "note": "5.4~5.5 휴강"},
    {"week": "7주", "period": "5.11~5.15", "topic": "생성형 AI 미니프로젝트", "kind": "프로젝트", "note": "기타교과 포함"},
    {"week": "8주", "period": "5.18~5.22", "topic": "서비스 개발/제안 + 미니프로젝트", "kind": "프로젝트", "note": "AI/DX 트랙별 산출물"},
    {"week": "9주", "period": "5.25~5.29", "topic": "서비스 개발/제안 미니프로젝트", "kind": "프로젝트", "note": "대체휴일 포함"},
    {"week": "10주", "period": "6.1~6.5", "topic": "서비스 개발/제안", "kind": "수업", "note": "서비스 구현 또는 제안 전략"},
    {"week": "11주", "period": "6.8~6.12", "topic": "서비스 개발/제안 미니프로젝트", "kind": "프로젝트", "note": "산출물 개선"},
    {"week": "12주", "period": "6.15~6.19", "topic": "Cloud", "kind": "수업", "note": "IT Infra, Cloud Infra"},
    {"week": "13주", "period": "6.22~6.26", "topic": "Cloud + Cloud 미니프로젝트", "kind": "프로젝트", "note": "AIVLE DAY 2차"},
    {"week": "14~22주", "period": "6.29~8.28", "topic": "빅프로젝트", "kind": "프로젝트", "note": "주제 선정, 현직자 코칭, 구현, 발표"},
    {"week": "23주", "period": "8.31~9.4", "topic": "취업지원 + 수료식", "kind": "취업", "note": "모의면접, 취업플랫폼, 수료식"},
]

QUICK_QUESTION_GROUPS = {
    "과정 이해": [
        "AI 개발자 트랙과 DX 컨설턴트 트랙의 차이를 정리해줘",
        "에이블스쿨 9기 전체 학습 흐름을 요약해줘",
        "비전공자가 먼저 확인해야 할 내용을 알려줘",
    ],
    "학습 전략": [
        "분석형 AI 예습은 무엇부터 하면 돼?",
        "생성형 AI 수업 전에 알아야 할 핵심 개념을 정리해줘",
        "Cloud 파트에서 어려워질 수 있는 부분을 알려줘",
    ],
    "프로젝트": [
        "미니프로젝트를 시작할 때 가장 먼저 해야 할 일을 알려줘",
        "빅프로젝트 주제 선정 기준을 정리해줘",
        "포트폴리오에 프로젝트를 정리할 때 주의할 점을 알려줘",
    ],
    "운영·일정": [
        "출결 기준에서 지각과 조퇴 기준을 알려줘",
        "수료 기준을 백서 기준으로 설명해줘",
        "에이블에듀에서 자주 쓰는 메뉴를 정리해줘",
    ],
}

LEVEL_GUIDE = {
    "초급": ["핵심 용어를 다시 정리", "예습 자료를 10분 단위로 나누어 재학습", "오답 문항을 같은 날 다시 풀기"],
    "중급": ["틀린 주제 중심으로 짧은 실습 진행", "개념을 미니프로젝트 산출물과 연결", "스터디에서 설명자 역할 수행"],
    "고급": ["심화 자료와 공식 문서 확인", "동료 질문 답변으로 설명력 강화", "프로젝트 적용 아이디어 도출"],
}

st.set_page_config(page_title=APP_TITLE, page_icon="📘", layout="wide", initial_sidebar_state="expanded")


# ============================================================
# 디자인: 밝은 학습 플랫폼 UI + Pretendard
# ============================================================

st.markdown(textwrap.dedent(
    """
    <style>
    @import url('https://cdn.jsdelivr.net/gh/orioncactus/pretendard/dist/web/variable/pretendardvariable-dynamic-subset.css');

    :root {
        --primary: #2F6F5E;
        --primary-hover: #255C4E;
        --secondary: #DFAE43;
        --accent: #6C7A89;
        --background: #F7F5EF;
        --surface: #FFFFFF;
        --surface-soft: #F1EEE7;
        --text: #1F2933;
        --muted: #68717A;
        --border: #DDD6C8;
        --success: #3F8F68;
        --warning: #C9822B;
        --error: #B94A48;
        --info: #4F7C8A;
        --shadow: 0 12px 28px rgba(31, 41, 51, .07);
        --shadow-soft: 0 6px 16px rgba(31, 41, 51, .05);
    }

    html, body, [class*="css"], .stApp, .stMarkdown, .stTextInput, .stTextArea,
    .stSelectbox, .stButton, .stDataFrame, .stChatMessage, input, textarea, button,
    .stRadio, .stCheckbox, .stFileUploader, .stTabs {
        font-family: 'Pretendard Variable', Pretendard, -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif !important;
    }

    .stApp {
        background:
            radial-gradient(circle at 3% 0%, rgba(47,111,94,.10), transparent 28%),
            radial-gradient(circle at 95% 3%, rgba(223,174,67,.12), transparent 24%),
            linear-gradient(180deg, #FBFAF6 0%, var(--background) 58%, #F1EEE7 100%) !important;
        color: var(--text);
    }

    .main .block-container {
        max-width: 1240px;
        padding-top: 1.1rem;
        padding-bottom: 3rem;
    }

    h1, h2, h3, h4, h5, h6,
    .stMarkdown h1, .stMarkdown h2, .stMarkdown h3, .stMarkdown h4 {
        color: #111827 !important;
        letter-spacing: -0.035em;
        font-weight: 880 !important;
    }
    p, li, label, .stMarkdown, .stCaption, .stText { color: var(--text); }
    small, .small-help { color: var(--muted) !important; font-size: .82rem; line-height: 1.5; }

    header[data-testid="stHeader"] { background: transparent !important; }
    [data-testid="stToolbar"] {
        display: flex !important;
        visibility: visible !important;
        height: auto !important;
        pointer-events: auto !important;
    }
    [data-testid="stStatusWidget"], [data-testid="stDeployButton"], .stDeployButton {
        display: none !important;
        visibility: hidden !important;
        height: 0 !important;
    }

    section[data-testid="stSidebar"] {
        background: linear-gradient(180deg, #FFFEFB 0%, var(--surface-soft) 100%) !important;
        border-right: 1px solid var(--border);
        min-width: 18rem;
    }
    section[data-testid="stSidebar"] * { color: var(--text) !important; }
    section[data-testid="stSidebar"] [data-testid="stMarkdownContainer"] p { color: var(--text) !important; }
    section[data-testid="stSidebar"] div[role="radiogroup"] label {
        border-radius: 14px !important;
        padding: 7px 8px !important;
        margin: 2px 0 !important;
    }
    section[data-testid="stSidebar"] div[role="radiogroup"] label:hover {
        background: rgba(47,111,94,.08) !important;
    }

    .app-shell { width: 100%; }
    .learning-hero {
        background: linear-gradient(135deg, #FFFFFF 0%, #F7F5EF 58%, #F1EEE7 100%);
        border: 1px solid var(--border);
        border-radius: 26px;
        box-shadow: var(--shadow);
        padding: 28px 32px;
        margin-bottom: 18px;
        position: relative;
        overflow: hidden;
    }
    .learning-hero::after {
        content: "";
        position: absolute;
        right: -72px;
        top: -78px;
        width: 220px;
        height: 220px;
        border-radius: 999px;
        background: rgba(47,111,94,.10);
    }
    .learning-hero::before {
        content: "";
        position: absolute;
        right: 62px;
        bottom: -78px;
        width: 160px;
        height: 160px;
        border-radius: 999px;
        background: rgba(223,174,67,.12);
    }
    .hero-kicker {
        display: inline-flex;
        align-items: center;
        gap: 7px;
        color: var(--primary);
        background: rgba(47,111,94,.09);
        border: 1px solid rgba(47,111,94,.20);
        padding: 6px 12px;
        border-radius: 999px;
        font-size: .82rem;
        font-weight: 850;
        margin-bottom: 10px;
    }
    .learning-hero h1 {
        margin: 0 0 8px 0;
        font-size: clamp(1.55rem, 3vw, 2.15rem);
        line-height: 1.2;
        color: var(--text) !important;
    }
    .learning-hero p {
        margin: 0;
        color: var(--muted);
        font-size: 1.0rem;
        line-height: 1.65;
        max-width: 860px;
    }

    .section-title {
        color: #111827 !important;
        font-size: 1.2rem;
        font-weight: 900;
        margin: 20px 0 8px 0;
        letter-spacing: -0.03em;
    }
    .section-sub {
        color: var(--muted);
        margin-top: -3px;
        margin-bottom: 12px;
        font-size: .93rem;
    }

    .summary-grid {
        display: grid;
        grid-template-columns: repeat(4, minmax(0, 1fr));
        gap: 12px;
        margin: 14px 0 18px 0;
    }
    .summary-card, .action-card, .path-card, .progress-card, .notice-card, .empty-state, .source-card, .info-panel, .learn-card {
        background: var(--surface);
        border: 1px solid var(--border);
        border-radius: 22px;
        box-shadow: var(--shadow-soft);
    }
    .summary-card { padding: 16px 17px; min-height: 112px; }
    .summary-card .label { color: var(--muted); font-size: .78rem; font-weight: 850; margin-bottom: 8px; }
    .summary-card .value { color: var(--text); font-weight: 920; font-size: 1.22rem; line-height: 1.25; word-break: break-word; }
    .summary-card .hint { color: var(--muted); font-size: .78rem; margin-top: 7px; }

    div[data-testid="stMetric"] {
        background: var(--surface) !important;
        border: 1px solid var(--border) !important;
        border-radius: 20px !important;
        box-shadow: var(--shadow-soft) !important;
        padding: 14px 16px !important;
    }
    div[data-testid="stMetric"] * { color: #111827 !important; }

    .learn-card, .info-panel, .action-card, .path-card, .progress-card { padding: 19px; margin-bottom: 14px; }
    .learn-card { min-height: 150px; transition: all .15s ease; }
    .learn-card:hover, .action-card:hover, .path-card:hover {
        transform: translateY(-1px);
        box-shadow: var(--shadow);
        border-color: rgba(47,111,94,.34);
    }
    .learn-card .badge, .status-badge {
        display: inline-flex;
        align-items: center;
        gap: 6px;
        padding: 5px 10px;
        border-radius: 999px;
        font-weight: 850;
        font-size: .78rem;
        color: var(--primary);
        background: rgba(47,111,94,.09);
        border: 1px solid rgba(47,111,94,.18);
        margin-bottom: 10px;
    }
    .learn-card h3, .action-card h3, .path-card h3 { font-size: 1.06rem; margin: 0 0 8px 0; color: #111827 !important; }
    .learn-card p, .action-card p, .path-card p, .progress-card p { margin: 0; color: var(--muted); line-height: 1.55; font-size: .94rem; }

    .routine-strip {
        display: grid;
        grid-template-columns: repeat(4, minmax(0, 1fr));
        gap: 12px;
        margin: 12px 0 20px 0;
    }
    .routine-item, .path-card {
        background: var(--surface);
        border: 1px solid var(--border);
        border-radius: 20px;
        padding: 15px;
        min-height: 112px;
        box-shadow: var(--shadow-soft);
    }
    .routine-item b { display:block; color: #111827; margin-bottom: 6px; }
    .routine-item span { color: var(--primary); font-size: .78rem; font-weight: 900; }
    .routine-item p { color: var(--muted); font-size: .88rem; line-height: 1.45; margin: 6px 0 0 0; }

    .notice-card {
        border-left: 5px solid var(--info);
        padding: 14px 16px;
        margin: 10px 0 14px 0;
        color: var(--text);
        background: #FBFAF6;
    }
    .notice-card b { color: #111827; }
    .empty-state {
        padding: 22px;
        background: #FFFEFB;
        text-align: left;
        margin: 12px 0;
    }
    .empty-state b { color:#111827; }
    .empty-state p { color:var(--muted); margin:.35rem 0 0; }
    .source-card, .source-box {
        background: #FBFAF6;
        border: 1px dashed var(--border);
        border-radius: 16px;
        padding: 14px;
        margin: 8px 0;
    }
    .source-box b, .source-card b { color: #111827; }
    .source-box p, .source-card p { color: var(--muted); }

    .badge-success { background: rgba(63,143,104,.12); color: var(--success); border-color: rgba(63,143,104,.25); }
    .badge-warning { background: rgba(201,130,43,.12); color: var(--warning); border-color: rgba(201,130,43,.25); }
    .badge-info { background: rgba(79,124,138,.12); color: var(--info); border-color: rgba(79,124,138,.25); }
    .badge-muted { background: var(--surface-soft); color: var(--muted); border-color: var(--border); }

    .sidebar-brand {
        display: flex;
        align-items: center;
        gap: 12px;
        border-radius: 22px;
        padding: 15px 13px;
        background: linear-gradient(135deg, #FFFFFF, #F1EEE7);
        border: 1px solid var(--border);
        box-shadow: var(--shadow-soft);
        margin: 4px 0 14px 0;
    }
    .brand-mark {
        width: 43px;
        height: 43px;
        display: grid;
        place-items: center;
        border-radius: 15px;
        color: white !important;
        background: var(--primary);
        font-weight: 950;
    }
    .brand-title { color: var(--text) !important; font-size: 1.05rem; font-weight: 950; line-height: 1.2; }
    .brand-sub { color: var(--primary) !important; font-size: .8rem; font-weight: 800; margin-top: 2px; }
    .current-pill {
        border-radius: 999px;
        padding: 9px 12px;
        background: rgba(47,111,94,.09);
        border: 1px solid rgba(47,111,94,.20);
        color: var(--primary) !important;
        font-weight: 850;
        font-size: .86rem;
        margin: 8px 0 10px 0;
    }
    .side-label {
        color: var(--muted) !important;
        font-size: .80rem;
        font-weight: 900;
        letter-spacing: .02em;
        margin: 12px 0 6px 0;
    }
    .sidebar-section {
        border-top: 1px solid var(--border);
        padding-top: 10px;
        margin-top: 12px;
    }

    div.stButton > button {
        border-radius: 14px !important;
        border: 1px solid rgba(47,111,94,.22) !important;
        background: linear-gradient(180deg, var(--primary), var(--primary-hover)) !important;
        color: #ffffff !important;
        font-weight: 850 !important;
        min-height: 2.75rem;
        box-shadow: 0 8px 18px rgba(47,111,94,.16);
    }
    div.stButton > button:hover {
        transform: translateY(-1px);
        border-color: rgba(47,111,94,.42) !important;
        box-shadow: 0 12px 24px rgba(47,111,94,.22);
    }
    div.stButton > button:disabled {
        background: #D8D2C7 !important;
        color: #7A756C !important;
        border-color: var(--border) !important;
        box-shadow: none !important;
    }

    .stTabs [data-baseweb="tab-list"] { gap: 8px; flex-wrap: wrap; }
    .stTabs [data-baseweb="tab"] {
        background: var(--surface);
        border: 1px solid var(--border);
        border-radius: 999px;
        padding: 8px 14px;
        color: var(--text);
    }
    .stTabs [aria-selected="true"] {
        background: rgba(47,111,94,.10) !important;
        color: var(--primary) !important;
        border-color: rgba(47,111,94,.30) !important;
    }

    [data-testid="stChatMessage"] {
        background: var(--surface);
        border: 1px solid var(--border);
        border-radius: 20px;
        padding: 10px 12px;
        box-shadow: var(--shadow-soft);
    }
    [data-testid="stDataFrame"] { border-radius: 18px; overflow: hidden; }
    .danger-zone { border-left: 5px solid var(--error); background: #FFF7F6; }


    .chat-shell {
        background: var(--surface);
        border: 1px solid var(--border);
        border-radius: 24px;
        box-shadow: var(--shadow-soft);
        padding: 14px 16px;
        margin: 10px 0 14px 0;
    }
    .chat-toolbar {
        display: flex;
        align-items: center;
        justify-content: space-between;
        gap: 12px;
        padding-bottom: 10px;
        border-bottom: 1px solid var(--border);
        margin-bottom: 12px;
    }
    .chat-title { font-weight: 900; color: #111827; }
    .chat-hint { color: var(--muted); font-size: .84rem; }
    .chat-window-note {
        color: var(--muted);
        font-size: .84rem;
        margin: -2px 0 8px 0;
    }


    [data-testid="collapsedControl"], [data-testid="stSidebarCollapsedControl"],
    button[data-testid="stExpandSidebarButton"], [data-testid="stExpandSidebarButton"],
    button[data-testid="stCollapseSidebarButton"], [data-testid="stCollapseSidebarButton"] {
        display: flex !important;
        visibility: visible !important;
        opacity: 1 !important;
        pointer-events: auto !important;
        z-index: 999999 !important;
        position: fixed !important;
        top: .75rem !important;
        left: .75rem !important;
        width: 2.65rem !important;
        height: 2.65rem !important;
        align-items: center !important;
        justify-content: center !important;
        background: #FFFFFF !important;
        border: 1px solid var(--border) !important;
        border-radius: 999px !important;
        box-shadow: var(--shadow) !important;
    }
    [data-testid="collapsedControl"] button,
    [data-testid="stSidebarCollapsedControl"] button,
    button[data-testid="stExpandSidebarButton"] *,
    [data-testid="stExpandSidebarButton"] *,
    button[data-testid="stCollapseSidebarButton"] *,
    [data-testid="stCollapseSidebarButton"] * {
        display: flex !important;
        visibility: visible !important;
        opacity: 1 !important;
        pointer-events: auto !important;
        color: var(--primary) !important;
        background: transparent !important;
    }
    button[data-testid="stBaseButton-headerNoPadding"][kind="headerNoPadding"] {
        visibility: visible !important;
        opacity: 1 !important;
        pointer-events: auto !important;
        z-index: 999999 !important;
        color: var(--primary) !important;
    }
    button[data-testid="stBaseButton-headerNoPadding"][kind="headerNoPadding"] * {
        visibility: visible !important;
        opacity: 1 !important;
        color: var(--primary) !important;
    }

    .calendar-fallback {
        width: 100%;
        border-collapse: separate;
        border-spacing: 8px;
        table-layout: fixed;
        margin-top: 10px;
    }
    .calendar-fallback th {
        color: #1F2933;
        font-weight: 900;
        padding: 8px;
        text-align: center;
    }
    .calendar-fallback td {
        vertical-align: top;
        background: #FFFFFF;
        border: 1px solid #DDD6C8;
        border-radius: 16px;
        min-height: 112px;
        height: 112px;
        padding: 10px;
        box-shadow: 0 6px 16px rgba(31, 41, 51, .04);
        overflow: hidden;
    }
    .calendar-fallback .muted-day { opacity: .35; background: #F1EEE7; }
    .calendar-fallback .day-num { font-weight: 900; color: #1F2933; margin-bottom: 6px; }
    .calendar-event-pill {
        display: block;
        margin: 3px 0;
        padding: 4px 7px;
        border-radius: 999px;
        background: rgba(47,111,94,.10);
        color: #2F6F5E;
        border: 1px solid rgba(47,111,94,.20);
        font-size: .76rem;
        white-space: nowrap;
        overflow: hidden;
        text-overflow: ellipsis;
    }
    .calendar-legend {
        display: flex;
        flex-wrap: wrap;
        gap: 8px;
        margin: 8px 0 12px 0;
    }
    .calendar-legend span {
        display: inline-flex;
        align-items: center;
        gap: 6px;
        border: 1px solid var(--border);
        border-radius: 999px;
        background: #FFFFFF;
        padding: 5px 10px;
        color: var(--text);
        font-size: .8rem;
        font-weight: 850;
    }
    .calendar-legend i {
        width: 10px;
        height: 10px;
        border-radius: 999px;
        display: inline-block;
    }

    @media (max-width: 1199px) { .summary-grid { grid-template-columns: repeat(2, minmax(0, 1fr)); } .routine-strip { grid-template-columns: repeat(2, minmax(0, 1fr)); } }
    @media (max-width: 767px) { .summary-grid, .routine-strip { grid-template-columns: 1fr; } .learning-hero { padding: 22px 18px; } .main .block-container { padding-left: .9rem; padding-right: .9rem; } }
    </style>
    """
).strip(), unsafe_allow_html=True)

# ============================================================
# 저장소 / 상태
# ============================================================

@dataclass
class SearchHit:
    index: int
    score: float
    title: str
    text: str


def safe_text(value: Any) -> str:
    """사용자 입력값이나 파일명 등이 HTML로 노출되지 않도록 이스케이프합니다."""
    return html.escape(str(value or ""), quote=True)


def html_block(markup: str) -> None:
    """들여쓰기 때문에 HTML이 코드블록으로 보이는 문제를 방지합니다."""
    st.markdown(textwrap.dedent(markup).strip(), unsafe_allow_html=True)


def read_json(path: Path, default: Any) -> Any:
    if not path.exists():
        return default
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return default


def write_json(path: Path, payload: Any) -> None:
    path.parent.mkdir(exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def get_config_value(name: str, default: Optional[str] = None) -> Optional[str]:
    value = None
    try:
        value = st.secrets.get(name)
    except Exception:
        value = None
    return value or os.getenv(name) or default


def get_api_key() -> Optional[str]:
    return get_config_value("OPENROUTER_API_KEY")


def get_login_credentials() -> Tuple[str, str]:
    return (
        get_config_value("APP_LOGIN_ID", "admin") or "admin",
        get_config_value("APP_LOGIN_PASSWORD", "aivle2026") or "aivle2026",
    )


def is_managed_whitepaper_path(path: Path) -> bool:
    try:
        resolved_parent = path.resolve().parent
        resolved_whitepaper_dir = WHITEPAPER_DIR.resolve()
    except Exception:
        return False
    return (
        resolved_parent == resolved_whitepaper_dir
        and path.stem == CURRENT_WHITEPAPER_STEM
        and path.suffix.lower() in WHITEPAPER_EXTENSIONS
    )


def current_uploaded_whitepaper_path() -> Optional[Path]:
    candidates = [
        path
        for path in WHITEPAPER_DIR.glob(f"{CURRENT_WHITEPAPER_STEM}.*")
        if path.is_file() and is_managed_whitepaper_path(path)
    ]
    if not candidates:
        return None
    return max(candidates, key=lambda path: path.stat().st_mtime)


def default_whitepaper_candidate() -> Path:
    return current_uploaded_whitepaper_path() or DEFAULT_WHITEPAPER_PATH


def get_whitepaper_meta() -> Dict[str, Any]:
    fallback_path = default_whitepaper_candidate()
    default = {
        "display_name": fallback_path.name,
        "storage_path": str(fallback_path),
        "updated_at": "",
        "size_bytes": fallback_path.stat().st_size if fallback_path.exists() else 0,
    }
    meta = read_json(WHITEPAPER_META_PATH, default)
    if not isinstance(meta, dict):
        return default
    merged = {**default, **meta}
    meta_path = Path(str(merged.get("storage_path", "")))
    if not (meta_path.exists() and is_managed_whitepaper_path(meta_path)):
        merged["storage_path"] = str(fallback_path)
        merged["display_name"] = fallback_path.name
        merged["size_bytes"] = fallback_path.stat().st_size if fallback_path.exists() else 0
    return merged


def get_whitepaper_path() -> Path:
    meta = get_whitepaper_meta()
    path = Path(str(meta.get("storage_path", DEFAULT_WHITEPAPER_PATH)))
    if path.exists() and (is_managed_whitepaper_path(path) or path == DEFAULT_WHITEPAPER_PATH):
        return path
    return default_whitepaper_candidate()


def save_whitepaper_meta(display_name: str, path: Path, size_bytes: int) -> None:
    if not is_managed_whitepaper_path(path):
        raise ValueError("Whitepaper uploads must be stored as data/whitepaper/current_whitepaper.*")
    write_json(
        WHITEPAPER_META_PATH,
        {
            "display_name": display_name,
            "storage_path": str(path),
            "updated_at": datetime.now().isoformat(timespec="seconds"),
            "size_bytes": size_bytes,
        },
    )



def get_auth_query_token() -> str:
    """URL query parameter에서 로그인 유지 토큰을 안전하게 읽습니다."""
    try:
        value = st.query_params.get("auth", "")
    except Exception:
        try:
            value = st.experimental_get_query_params().get("auth", [""])
        except Exception:
            value = ""
    if isinstance(value, list):
        value = value[0] if value else ""
    return str(value or "")


def auth_secret() -> str:
    """토큰 서명용 문자열입니다. 별도 secret이 없으면 로그인 비밀번호를 사용합니다."""
    return get_config_value("APP_LOGIN_SECRET") or get_config_value("APP_LOGIN_PASSWORD", "aivle2026") or "aivle2026"


def make_auth_token(login_id: str) -> str:
    payload = str(login_id or "")
    signature = hmac.new(auth_secret().encode("utf-8"), payload.encode("utf-8"), hashlib.sha256).hexdigest()
    return f"{payload}.{signature}"


def is_valid_auth_token(token: str) -> bool:
    if not token or "." not in token:
        return False
    login_id, signature = token.rsplit(".", 1)
    expected_id, _ = get_login_credentials()
    if not hmac.compare_digest(login_id, expected_id):
        return False
    expected = make_auth_token(login_id).rsplit(".", 1)[1]
    return hmac.compare_digest(signature, expected)


def set_auth_query_token(login_id: str) -> None:
    try:
        st.query_params["auth"] = make_auth_token(login_id)
    except Exception:
        try:
            st.experimental_set_query_params(auth=make_auth_token(login_id))
        except Exception:
            pass


def clear_auth_query_token() -> None:
    try:
        if "auth" in st.query_params:
            del st.query_params["auth"]
    except Exception:
        try:
            st.experimental_set_query_params()
        except Exception:
            pass


def sync_sidebar_page() -> None:
    selected = st.session_state.get("sidebar_page", "대시보드")
    st.session_state["active_page"] = selected if selected in MENU_OPTIONS else "대시보드"

def init_state() -> None:
    st.session_state.setdefault("authenticated", False)
    st.session_state.setdefault("login_error", "")
    st.session_state.setdefault("active_page", "대시보드")
    st.session_state.setdefault("sidebar_page", st.session_state.get("active_page", "대시보드"))
    st.session_state.setdefault("nav_target", None)
    st.session_state.setdefault("last_whitepaper_upload_sig", "")

    if not st.session_state.get("authenticated") and is_valid_auth_token(get_auth_query_token()):
        st.session_state["authenticated"] = True
        st.session_state["login_error"] = ""

    if st.session_state["active_page"] not in MENU_OPTIONS:
        st.session_state["active_page"] = "대시보드"
    if st.session_state["sidebar_page"] not in MENU_OPTIONS:
        st.session_state["sidebar_page"] = st.session_state["active_page"]


def request_nav(page: str) -> None:
    st.session_state["nav_target"] = page if page in MENU_OPTIONS else "대시보드"
    st.rerun()


def authenticate(login_id: str, login_password: str) -> bool:
    expected_id, expected_pw = get_login_credentials()
    return hmac.compare_digest(login_id, expected_id) and hmac.compare_digest(login_password, expected_pw)


# ============================================================
# 파일 파싱
# ============================================================


def parse_docx_path(path: Path) -> str:
    if Document is None:
        return ""
    doc = Document(str(path))
    lines: List[str] = []
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            text = Paragraph(child, doc).text.strip()
            if text:
                lines.append(text)
        elif isinstance(child, CT_Tbl):
            table = Table(child, doc)
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                cells = [cell for cell in cells if cell]
                if cells:
                    lines.append(" | ".join(cells))
    return "\n".join(lines)


def parse_pdf_path(path: Path) -> str:
    if PdfReader is None:
        return ""
    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def parse_text_file(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return parse_docx_path(path)
    if suffix == ".pdf":
        return parse_pdf_path(path)
    return path.read_text(encoding="utf-8", errors="ignore")


def parse_uploaded_file(uploaded: Any, max_chars: int = 60000) -> str:
    if uploaded is None:
        return ""
    name = uploaded.name
    suffix = Path(name).suffix.lower()
    data = uploaded.getvalue()
    try:
        if suffix == ".pdf":
            if PdfReader is None:
                return ""
            reader = PdfReader(BytesIO(data))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        elif suffix == ".docx":
            if Document is None:
                return ""
            doc = Document(BytesIO(data))
            lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip().replace("\n", " ") for cell in row.cells if cell.text.strip()]
                    if cells:
                        lines.append(" | ".join(cells))
            text = "\n".join(lines)
        elif suffix in [".xlsx", ".xls"]:
            sheets = pd.read_excel(BytesIO(data), sheet_name=None)
            blocks = []
            for sheet_name, df in sheets.items():
                blocks.append(f"[시트: {sheet_name}]")
                blocks.append(df.fillna("").astype(str).to_csv(index=False))
            text = "\n".join(blocks)
        elif suffix == ".csv":
            text = data.decode("utf-8-sig", errors="ignore")
        else:
            text = data.decode("utf-8", errors="ignore")
    except Exception:
        return ""
    parsed = normalize_text(text)[:max_chars]
    return parsed


def normalize_text(text: str) -> str:
    text = text.replace("\r", "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def split_into_chunks(text: str, size: int = 650, overlap: int = 80) -> List[Dict[str, str]]:
    clean = normalize_text(text)
    if not clean:
        return []
    paragraphs = [p.strip() for p in re.split(r"\n+", clean) if p.strip()]
    chunks: List[Dict[str, str]] = []
    title = "백서"
    buf = ""
    for para in paragraphs:
        if re.match(r"^(Ⅰ|Ⅱ|Ⅲ|Ⅳ|Ⅴ|Ⅵ|Ⅶ|Ⅷ|Ⅸ|Ⅹ)\.", para) or re.match(r"^\d+\.\s", para):
            title = para[:70]
        if len(buf) + len(para) + 1 <= size:
            buf = f"{buf}\n{para}".strip()
        else:
            if buf:
                chunks.append({"title": title, "text": buf})
            tail = buf[-overlap:] if overlap and buf else ""
            buf = f"{tail}\n{para}".strip()
    if buf:
        chunks.append({"title": title, "text": buf})
    return chunks


@st.cache_resource(show_spinner=False)
def build_index(path_string: str, modified_at: float, file_size: int) -> Dict[str, Any]:
    path = Path(path_string)
    raw_text = parse_text_file(path)
    chunks = split_into_chunks(raw_text)
    if not chunks:
        return {"text": raw_text, "chunks": [], "vectorizer": None, "matrix": None}
    texts = [chunk["text"] for chunk in chunks]
    try:
        vectorizer = TfidfVectorizer(analyzer="char_wb", ngram_range=(2, 5), max_features=30000)
        matrix = vectorizer.fit_transform(texts)
    except ValueError:
        return {"text": raw_text, "chunks": chunks, "vectorizer": None, "matrix": None}
    return {"text": raw_text, "chunks": chunks, "vectorizer": vectorizer, "matrix": matrix}


def load_index() -> Optional[Dict[str, Any]]:
    path = get_whitepaper_path()
    if not path.exists():
        return None
    stat = path.stat()
    return build_index(str(path), stat.st_mtime, stat.st_size)


def search_whitepaper(query: str, k: int = 6) -> List[SearchHit]:
    index = load_index()
    if not index or not index.get("chunks"):
        return []
    chunks = index["chunks"]
    vectorizer = index.get("vectorizer")
    matrix = index.get("matrix")
    if vectorizer is None or matrix is None:
        return [SearchHit(i, 0.0, chunk["title"], chunk["text"]) for i, chunk in enumerate(chunks[:k])]
    query_vector = vectorizer.transform([query])
    base_scores = cosine_similarity(query_vector, matrix).ravel()
    tokens = [t for t in re.findall(r"[가-힣A-Za-z0-9]+", query) if len(t) >= 2]
    adjusted: List[Tuple[float, int]] = []
    for idx, chunk in enumerate(chunks):
        blob = f"{chunk['title']} {chunk['text']}".lower()
        token_hits = sum(1 for token in tokens if token.lower() in blob)
        exact_bonus = 0.10 if query.replace(" ", "") in blob.replace(" ", "") else 0.0
        title_bonus = sum(1 for token in tokens if token.lower() in chunk['title'].lower()) * 0.035
        score = float(base_scores[idx]) + token_hits * 0.035 + exact_bonus + title_bonus
        adjusted.append((score, idx))
    ranked = sorted(adjusted, key=lambda item: item[0], reverse=True)[:k]
    hits = []
    for score, idx in ranked:
        if score <= 0 and hits:
            continue
        chunk = chunks[idx]
        hits.append(SearchHit(index=idx, score=float(score), title=chunk["title"], text=chunk["text"]))
    return hits


# ============================================================
# LLM / 생성 기능
# ============================================================


class OpenRouterCallError(RuntimeError):
    """API key를 포함하지 않는 안전한 OpenRouter 오류입니다."""


def safe_openrouter_error(error: Exception) -> str:
    if isinstance(error, OpenRouterCallError) and str(error):
        return str(error)
    response = getattr(error, "response", None)
    status_code = getattr(response, "status_code", None)
    if status_code:
        return f"HTTP {status_code}"
    return error.__class__.__name__


def call_openrouter_model(
    model: str,
    messages: List[Dict[str, str]],
    api_key: str,
    temperature: float,
    max_tokens: int,
) -> str:
    payload = {
        "model": model,
        "messages": messages,
        "temperature": temperature,
        "max_tokens": max_tokens,
    }
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        # HTTP header values must be latin-1 encodable. Keep this ASCII-only.
        "X-Title": OPENROUTER_APP_TITLE,
    }
    try:
        response = requests.post(
            OPENROUTER_ENDPOINT,
            headers=headers,
            json=payload,
            timeout=OPENROUTER_TIMEOUT_SECONDS,
        )
        response.raise_for_status()
    except requests.RequestException as error:
        raise OpenRouterCallError(safe_openrouter_error(error)) from error

    try:
        data = response.json()
    except ValueError as error:
        raise OpenRouterCallError("invalid_json_response") from error

    try:
        content = data["choices"][0]["message"]["content"]
    except (KeyError, IndexError, TypeError) as error:
        raise OpenRouterCallError("missing_message_content") from error

    if not isinstance(content, str) or not content.strip():
        raise OpenRouterCallError("empty_message_content")
    return content.strip()


def call_openrouter(
    messages: List[Dict[str, str]],
    temperature: float = 0.2,
    max_tokens: int = 1000,
) -> Optional[str]:
    """
    Nvidia OpenRouter 모델을 단일 메인 모델로 호출합니다.
    기존 호출부 호환을 위해 함수명은 유지합니다.
    """

    api_key = get_api_key()
    if not api_key:
        return None

    try:
        return call_openrouter_model(
            DEFAULT_OPENROUTER_MODEL,
            messages,
            api_key,
            temperature,
            max_tokens,
        )
    except OpenRouterCallError as error:
        logger.error("OpenRouter Nvidia model call failed. reason=%s", safe_openrouter_error(error))
        st.error("현재 답변 생성에 실패했습니다.")
        return None


def ask_llm(system_prompt: str, user_prompt: str, temperature: float = 0.15) -> Optional[str]:
    return call_openrouter(
        [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=temperature,
        max_tokens=1800,
    )



# ============================================================
# RAG 검색 고도화: Query Rewriting + Multi Query + Sub Query
# - HyDE는 의도적으로 제외합니다.
# - API 키가 없어도 fallback 쿼리 확장으로 앱이 중단되지 않습니다.
# ============================================================

def extract_json_object(text: str) -> Optional[Dict[str, Any]]:
    """LLM 응답에서 JSON 객체만 안전하게 추출합니다."""
    if not text:
        return None
    candidates = re.findall(r"```json\s*(.*?)\s*```", text, flags=re.S)
    candidates.append(text)
    for candidate in candidates:
        start = candidate.find("{")
        end = candidate.rfind("}")
        if start == -1 or end == -1 or start >= end:
            continue
        try:
            parsed = json.loads(candidate[start : end + 1])
        except Exception:
            continue
        if isinstance(parsed, dict):
            return parsed
    return None


def unique_keep_order(items: List[str], limit: Optional[int] = None) -> List[str]:
    seen = set()
    result: List[str] = []
    for item in items:
        clean = normalize_text(str(item or ""))
        if not clean:
            continue
        key = clean.lower().replace(" ", "")
        if key in seen:
            continue
        seen.add(key)
        result.append(clean)
        if limit is not None and len(result) >= limit:
            break
    return result


def likely_complex_question(question: str) -> bool:
    q = str(question or "")
    markers = ["그리고", "또", "또는", "및", "까지", "각각", "차이", "비교", ",", "/", " vs ", "VS"]
    return any(marker in q for marker in markers)


def fallback_query_rewrite(question: str) -> str:
    """LLM을 사용할 수 없을 때 적용하는 안전한 재작성 쿼리입니다."""
    q = normalize_text(question)
    alias_terms: List[str] = []
    lowered = q.lower()
    for topic, aliases in TOPIC_ALIASES.items():
        topic_blob = " ".join([topic] + aliases).lower()
        if any(term.lower() in lowered for term in [topic] + aliases) or any(tok in topic_blob for tok in lowered.split()):
            alias_terms.extend([topic] + aliases[:5])
    if official_notice_needed(q):
        alias_terms.extend(["공식 기준", "최신 공지", "운영자 확인", "기수별 변동 가능"])
    if any(word in q for word in ["수료", "출석", "출결", "결석", "지각", "조퇴", "공가"]):
        alias_terms.extend(["출결", "지각", "조퇴", "외출", "공가", "결석", "수료", "소정훈련일수"])
    if any(word in q for word in ["포트폴리오", "프로젝트", "빅프로젝트", "미니프로젝트"]):
        alias_terms.extend(["프로젝트", "빅프로젝트", "산출물", "공개 제한", "GitHub", "저작권", "취업"])
    suffix = " ".join(unique_keep_order(alias_terms, limit=10))
    return normalize_text(f"{q} {suffix} 백서 기준 공식 안내")


def fallback_multi_queries(question: str, rewritten_query: str) -> List[str]:
    q = normalize_text(question)
    if any(word in q for word in ["수료", "출석", "출결", "결석", "지각", "조퇴", "공가"]):
        candidates = [
            "출결 기준 지각 조퇴 외출 결석 공가 수료 소정훈련일수",
            "정상 수료 전체 소정훈련일수 80% 이상 출석 결석일수 제적 기준",
            "에이블스쿨 출석인정 공가 증빙서류 지각 조퇴 외출 처리",
        ]
    elif any(word in q for word in ["트랙", "AI 개발자", "DX", "컨설턴트"]):
        candidates = [
            "AI 개발자 트랙 DX 컨설턴트 트랙 차이 인재상 산출물",
            "트랙별 커리큘럼 AI Cloud 서비스 개발 서비스 제안",
            "AI DX 공통 역량 데이터 AI Cloud 협업 발표 문서화",
        ]
    elif any(word in q for word in ["포트폴리오", "면접", "취업", "공고", "기업"]):
        candidates = [
            "포트폴리오 빅프로젝트 산출물 공개 제한 저작권 GitHub 블로그",
            "취업지원 잡페어 포트폴리오 면접 자기소개서 프로젝트 경험",
            "프로젝트 역할 기술 성과 문제정의 면접 질문 준비",
        ]
    else:
        candidates = [
            rewritten_query,
            f"{q} 핵심 개념 학습 목표 커리큘럼",
            f"{q} 예습 복습 프로젝트 적용 백서 근거",
        ]
    return unique_keep_order(candidates, limit=3)


def fallback_sub_queries(question: str) -> List[str]:
    q = normalize_text(question)
    if not likely_complex_question(q):
        return []
    protected = q.replace("AI 개발자 트랙과 DX 컨설턴트 트랙", "AI 개발자 트랙 / DX 컨설턴트 트랙")
    parts = re.split(r"\s*(?:그리고|또는|또|및|까지|,|/| vs | VS )\s*", protected)
    cleaned = []
    for part in parts:
        part = normalize_text(part)
        if len(part) >= 4:
            cleaned.append(part)
    if len(cleaned) >= 2:
        return unique_keep_order(cleaned, limit=4)
    if any(word in q for word in ["차이", "비교", "각각"]):
        return unique_keep_order([f"{q}의 첫 번째 항목", f"{q}의 두 번째 항목", f"{q}의 비교 기준"], limit=3)
    return []


def build_rag_query_plan(question: str) -> Dict[str, Any]:
    """
    백서 검색용 쿼리 계획을 생성합니다.
    포함: Query Rewriting, Multi Query, 조건부 Sub Query
    제외: HyDE 가상답변 생성
    """
    q = normalize_text(question)
    fallback_rewritten = fallback_query_rewrite(q)
    fallback_plan = {
        "rewritten_query": fallback_rewritten,
        "multi_queries": fallback_multi_queries(q, fallback_rewritten),
        "sub_queries": fallback_sub_queries(q),
    }

    system = (
        "너는 RAG 검색 쿼리 설계자입니다. 사용자의 질문을 백서 검색에 적합하게 정리합니다. "
        "가상 답변을 만들지 마세요. HyDE 방식은 사용하지 마세요. "
        "반드시 JSON 객체만 반환하세요."
    )
    user = f"""
[사용자 질문]
{q}

[작업]
1. rewritten_query: 백서 검색에 적합한 한 문장으로 재작성
2. multi_queries: 같은 의미를 다른 표현으로 찾기 위한 검색 쿼리 3개 이하
3. sub_queries: 질문이 복합 질문일 때만 하위 질문 4개 이하, 단순 질문이면 빈 배열

[JSON 형식]
{{
  "rewritten_query": "재작성 검색문",
  "multi_queries": ["검색문1", "검색문2", "검색문3"],
  "sub_queries": ["하위질문1", "하위질문2"]
}}
"""
    response = ask_llm(system, user, temperature=0.0)
    parsed = extract_json_object(response or "")
    if not isinstance(parsed, dict):
        return fallback_plan

    rewritten = normalize_text(str(parsed.get("rewritten_query", ""))) or fallback_rewritten
    multi_raw = parsed.get("multi_queries", [])
    sub_raw = parsed.get("sub_queries", [])
    multi = unique_keep_order([str(x) for x in multi_raw if isinstance(x, (str, int, float))], limit=3)
    sub = unique_keep_order([str(x) for x in sub_raw if isinstance(x, (str, int, float))], limit=4)
    if not multi:
        multi = fallback_plan["multi_queries"]
    if not likely_complex_question(q):
        sub = []
    elif not sub:
        sub = fallback_plan["sub_queries"]
    return {
        "rewritten_query": rewritten[:300],
        "multi_queries": multi,
        "sub_queries": sub,
    }


def advanced_search_whitepaper(question: str, k: int = 8) -> List[SearchHit]:
    """Query Rewriting + Multi Query + Sub Query를 적용한 통합 백서 검색입니다."""
    plan = build_rag_query_plan(question)
    weighted_queries: List[Tuple[str, float]] = [(question, 1.00)]
    if plan.get("rewritten_query"):
        weighted_queries.append((str(plan["rewritten_query"]), 0.98))
    weighted_queries.extend((q, 0.92) for q in plan.get("multi_queries", []))
    weighted_queries.extend((q, 0.88) for q in plan.get("sub_queries", []))

    deduped_queries: List[Tuple[str, float]] = []
    seen_queries = set()
    for query_text, weight in weighted_queries:
        clean_query = normalize_text(str(query_text))
        key = clean_query.lower().replace(" ", "")
        if not clean_query or key in seen_queries:
            continue
        seen_queries.add(key)
        deduped_queries.append((clean_query, weight))
        if len(deduped_queries) >= 10:
            break
    weighted_queries = deduped_queries

    aggregated: Dict[int, SearchHit] = {}
    for q_idx, (query_text, weight) in enumerate(weighted_queries):
        for hit in search_whitepaper(query_text, k=4):
            score = float(hit.score) * weight + max(0.0, 0.04 - q_idx * 0.004)
            existing = aggregated.get(hit.index)
            if existing is None or score > existing.score:
                aggregated[hit.index] = SearchHit(index=hit.index, score=score, title=hit.title, text=hit.text)

    hits = sorted(aggregated.values(), key=lambda item: item.score, reverse=True)[:k]
    if hits:
        return hits
    return search_whitepaper(question, k=k)


def make_context(hits: List[SearchHit]) -> str:
    return "\n\n---\n\n".join([f"[근거 {i}] 제목: {hit.title}\n내용:\n{hit.text}" for i, hit in enumerate(hits, start=1)])


def career_whitepaper_guide(query: str, k: int = 4) -> str:
    hits = search_whitepaper(query, k=k)
    return make_context(hits)


def official_notice_needed(question: str) -> bool:
    return any(keyword in question for keyword in OFFICIAL_NOTICE_KEYWORDS)


def fallback_answer(question: str, hits: List[SearchHit]) -> str:
    if not hits:
        return "백서에서 직접 확인할 수 있는 근거를 찾지 못했습니다. 질문 표현을 더 구체화하거나 학습 자료를 다시 업로드해 주세요."
    context = hits[0].text[:950]
    notice = "\n\n**주의사항**\n- 일정, 출결, 수료, 지원 자격처럼 변동 가능한 항목은 최신 공식 공지와 담당자 안내를 최종 기준으로 확인해야 합니다." if official_notice_needed(question) else ""
    return (
        "### 핵심 답변\n"
        f"백서에서 가장 관련성이 높은 내용은 다음입니다.\n\n{context}\n\n"
        "### 바로 할 일\n"
        "1. 위 내용에서 본인에게 해당하는 항목을 먼저 표시합니다.\n"
        "2. 모르는 용어를 학습 질의에 다시 질문합니다.\n"
        "3. 예습·진단 메뉴에서 같은 주제로 이해도를 확인합니다."
        f"{notice}"
    )


def answer_from_whitepaper(question: str) -> Tuple[str, List[SearchHit], List[Dict[str, str]]]:
    hits = advanced_search_whitepaper(question, k=8)
    context = make_context(hits)
    system = (
        "당신은 KT AIVLE School 학습자를 돕는 백서 기반 학습 도우미입니다. "
        "반드시 제공된 백서 근거 안에서만 답변하고, 근거가 부족하면 부족하다고 말합니다. "
        "내부 검색에는 Query Rewriting, Multi Query, Sub Query가 적용되지만 최종 답변은 실제 백서 근거만 사용합니다. "
        "일정, 출결, 수료, 지원 자격, 평가처럼 변동 가능한 내용은 최신 공식 공지 확인이 필요하다고 안내합니다. "
        "답변은 한국어 마크다운으로 작성합니다."
    )
    user = f"""
[백서 근거]
{context}

[질문]
{question}

[답변 형식]
### 핵심 답변
### 상세 설명
### 바로 할 일
### 주의사항
"""
    answer = ask_llm(system, user, temperature=0.1) or fallback_answer(question, hits)
    return answer, hits, link_recommendations(question)


def extract_links(text: str) -> List[str]:
    raw = re.findall(r"(?:https?://)?(?:[A-Za-z0-9-]+\.)+[A-Za-z]{2,}(?:/[A-Za-z0-9_./?=&%-]*)?", text)
    links: List[str] = []
    for item in raw:
        url = item if item.startswith("http") else "https://" + item
        if url not in links:
            links.append(url)
    return links


def link_recommendations(question: str) -> List[Dict[str, str]]:
    index = load_index()
    source_text = index.get("text", "") if index else ""
    base = [
        {"name": "KT AIVLE School 공식 홈페이지", "url": "https://aivle.kt.co.kr", "use": "모집·트랙·공식 안내 확인"},
        {"name": "AIVLE-EDU", "url": "https://aivle.edu.kt.co.kr", "use": "강의·출결·과제·커뮤니티 확인"},
        {"name": "고용24", "url": "https://www.work24.go.kr", "use": "K-DT 수강신청·국민내일배움카드 확인"},
    ]
    for url in extract_links(source_text):
        if not any(item["url"] == url for item in base):
            base.append({"name": url.replace("https://", ""), "url": url, "use": "백서에서 추출된 참고 링크"})
    tokens = re.findall(r"[가-힣A-Za-z0-9]+", question.lower())
    scored = []
    for row in base:
        blob = f"{row['name']} {row['url']} {row['use']}".lower()
        score = sum(1 for token in tokens if token and token in blob)
        scored.append((score, row))
    return [row for _, row in sorted(scored, key=lambda x: x[0], reverse=True)[:5]]


def generate_prep_material(week: str, topic: str, minutes: int) -> Tuple[str, List[SearchHit]]:
    query = f"{week} {topic} {' '.join(TOPIC_ALIASES.get(topic, []))}"
    hits = advanced_search_whitepaper(query, k=6)
    context = make_context(hits)
    system = "백서 근거로 수업 전 예습 자료를 만드는 학습 코치입니다. 근거 밖 추측은 하지 않습니다."
    user = f"""
[백서 근거]
{context}

[예습 조건]
- 주차: {week}
- 주제: {topic}
- 학습 가능 시간: {minutes}분

[출력 형식]
### 예습 목표
### 핵심 개념
### {minutes}분 학습 순서
### 확인 질문
### 수업 전 체크포인트
"""
    fallback = f"""### 예습 목표
{topic}의 핵심 용어와 학습 위치를 먼저 파악합니다.

### 핵심 개념
{hits[0].text[:800] if hits else '백서 근거를 찾지 못했습니다. 학습 자료를 확인해 주세요.'}

### {minutes}분 학습 순서
1. 관련 용어 5개 표시
2. 백서 근거 문단 읽기
3. 모르는 개념을 학습 질의에 질문
4. 확인 질문 3개 작성

### 확인 질문
- 이 주제가 전체 커리큘럼에서 어떤 위치인가?
- 수업 전에 알아야 할 용어는 무엇인가?
- 프로젝트와 어떻게 연결되는가?
"""
    return ask_llm(system, user, temperature=0.2) or fallback, hits


def fallback_quiz(topic: str) -> List[Dict[str, Any]]:
    aliases = TOPIC_ALIASES.get(topic, [topic])
    return [
        {"topic": topic, "question": f"{topic} 학습을 시작할 때 가장 먼저 해야 할 일은?", "choices": ["핵심 용어와 학습 목표 확인", "무작정 코드 복사", "공식 일정 무시", "시험만 먼저 보기"], "answer_index": 0, "explanation": "예습 단계에서는 학습 목표와 핵심 용어를 먼저 확인해야 합니다."},
        {"topic": topic, "question": "백서에 없는 내용은 어떻게 처리해야 하는가?", "choices": ["추측하지 않고 근거 부족으로 표시", "확정 정보처럼 말하기", "임의로 작성", "오래된 후기만 사용"], "answer_index": 0, "explanation": "백서 기반 앱은 근거가 없는 내용을 단정하지 않아야 합니다."},
        {"topic": topic, "question": f"{aliases[0]} 학습 후 이해도를 확인하는 데 적절한 기능은?", "choices": ["쪽지시험", "파일 삭제", "로그아웃", "테마 변경"], "answer_index": 0, "explanation": "쪽지시험은 예습 자료 기반 이해도 점검에 사용됩니다."},
        {"topic": topic, "question": "틀린 문제를 반복 복습하기 위한 기능은?", "choices": ["오답노트", "공고 정리", "로그인", "백서명 변경"], "answer_index": 0, "explanation": "오답노트는 틀린 문항과 해설을 누적해 복습하는 기능입니다."},
        {"topic": topic, "question": "출결·수료 같은 변동 가능 정보는 어떻게 확인해야 하는가?", "choices": ["최신 공식 공지 확인", "추측으로 판단", "친구 말만 기준", "오래된 자료만 사용"], "answer_index": 0, "explanation": "기수별로 달라질 수 있는 정보는 최신 공식 공지를 최종 기준으로 봐야 합니다."},
    ]


def extract_json_array(text: str) -> Optional[List[Dict[str, Any]]]:
    if not text:
        return None
    candidates = re.findall(r"```json\s*(.*?)\s*```", text, flags=re.S)
    candidates.append(text)
    for candidate in candidates:
        start = candidate.find("[")
        end = candidate.rfind("]")
        if start == -1 or end == -1 or start >= end:
            continue
        try:
            data = json.loads(candidate[start : end + 1])
        except Exception:
            continue
        if isinstance(data, list):
            return data
    return None


def normalize_quiz(data: Any, topic: str, count: int) -> List[Dict[str, Any]]:
    normalized: List[Dict[str, Any]] = []
    if not isinstance(data, list):
        return fallback_quiz(topic)[:count]
    for item in data:
        if not isinstance(item, dict):
            continue
        question = str(item.get("question", "")).strip()
        choices = item.get("choices")
        answer_index = item.get("answer_index", 0)
        explanation = str(item.get("explanation", "")).strip() or "백서 기반 핵심 개념을 확인합니다."
        if not question or not isinstance(choices, list) or len(choices) < 4:
            continue
        choices = [str(c).strip() for c in choices[:4]]
        try:
            answer_index = int(answer_index)
        except Exception:
            answer_index = 0
        if answer_index < 0 or answer_index >= len(choices):
            answer_index = 0
        normalized.append({"topic": str(item.get("topic", topic)), "question": question, "choices": choices, "answer_index": answer_index, "explanation": explanation})
    if len(normalized) < count:
        normalized.extend(fallback_quiz(topic))
    return normalized[:count]


def generate_quiz(topic: str, prep_material: str, count: int = 5) -> List[Dict[str, Any]]:
    system = "학습 자료를 기반으로 객관식 쪽지시험을 JSON 배열로만 생성합니다."
    user = f"""
[주제]
{topic}

[예습 자료]
{prep_material[:3500]}

[{count}문항 JSON 배열 형식]
[
  {{"topic":"주제", "question":"문제", "choices":["보기1","보기2","보기3","보기4"], "answer_index":0, "explanation":"해설"}}
]
"""
    response = ask_llm(system, user, temperature=0.2)
    parsed = extract_json_array(response or "")
    return normalize_quiz(parsed, topic, count)


def judge_level(score: float) -> str:
    if score >= 85:
        return "고급"
    if score >= 60:
        return "중급"
    return "초급"


def build_topic_stats(quiz: List[Dict[str, Any]], selected: Dict[int, int]) -> pd.DataFrame:
    rows: Dict[str, Dict[str, int]] = {}
    for i, item in enumerate(quiz):
        topic = item.get("topic", "기타")
        rows.setdefault(topic, {"문항수": 0, "정답수": 0})
        rows[topic]["문항수"] += 1
        if selected.get(i) == item.get("answer_index"):
            rows[topic]["정답수"] += 1
    result = []
    for topic, stat in rows.items():
        rate = round(stat["정답수"] / max(1, stat["문항수"]) * 100, 1)
        result.append({"주제": topic, "문항수": stat["문항수"], "정답수": stat["정답수"], "정답률": rate})
    return pd.DataFrame(result)


def save_result(payload: Dict[str, Any]) -> None:
    data = read_json(RESULT_PATH, [])
    if not isinstance(data, list):
        data = []
    data.append(payload)
    write_json(RESULT_PATH, data)


def save_wrong_notes(items: List[Dict[str, Any]]) -> None:
    data = read_json(WRONG_NOTE_PATH, [])
    if not isinstance(data, list):
        data = []
    data.extend(items)
    write_json(WRONG_NOTE_PATH, data)


def study_recommendation(level: str, weak_topics: List[str]) -> str:
    weak = ", ".join(weak_topics) if weak_topics else "현재 뚜렷한 취약 주제 없음"
    actions = LEVEL_GUIDE.get(level, LEVEL_GUIDE["초급"])
    return "\n".join(["### 수준별 스터디 추천", f"- 현재 등급: **{level}**", f"- 보완 주제: **{weak}**"] + [f"- {a}" for a in actions])


def summarize_notice(notice_text: str, interests: str) -> Dict[str, str]:
    guide = career_whitepaper_guide("취업지원 잡페어 포트폴리오 공모전 채용 프로젝트", k=3)
    system = "공고/공모전/채용형 프로그램 정보를 학습자 관점으로 정리합니다. JSON 객체만 반환합니다."
    user = f"""
관심 분야: {interests}

[백서 가이드 일부]
{guide or '백서 가이드 없음'}

공고 내용:
{notice_text[:3500]}

다음 JSON 형식으로만 반환:
{{"title":"공고명", "type":"공모전/프로젝트/채용/기타", "deadline":"YYYY-MM-DD 또는 미확인", "fit":"관심 분야와의 관련성", "actions":"다음 행동 2~3개"}}
"""
    response = ask_llm(system, user, temperature=0.1)
    if response:
        try:
            start = response.find("{")
            end = response.rfind("}")
            parsed = json.loads(response[start : end + 1])
            if isinstance(parsed, dict):
                return {"title": str(parsed.get("title", "미확인")), "type": str(parsed.get("type", "기타")), "deadline": str(parsed.get("deadline", "미확인")), "fit": str(parsed.get("fit", "미확인")), "actions": str(parsed.get("actions", "미확인"))}
        except Exception:
            pass
    date_match = re.search(r"(20\d{2})[.\-/년\s]+(\d{1,2})[.\-/월\s]+(\d{1,2})", notice_text)
    deadline = "미확인"
    if date_match:
        y, m, d = map(int, date_match.groups())
        deadline = f"{y:04d}-{m:02d}-{d:02d}"
    first_line = next((line.strip() for line in notice_text.splitlines() if line.strip()), "공고")
    return {"title": first_line[:60], "type": "기타", "deadline": deadline, "fit": f"관심 분야({interests})와 직접 비교가 필요합니다.", "actions": "지원 자격 확인, 마감일 등록, 제출물 목록 작성"}


def analyze_portfolio(portfolio_text: str, job_text: str, target_role: str) -> str:
    context = career_whitepaper_guide("포트폴리오 빅프로젝트 산출물 공개 제한 취업지원", k=5)
    if not portfolio_text.strip():
        return "포트폴리오 또는 프로젝트 정리표를 업로드해야 분석할 수 있습니다."
    system = (
        "AIVLE 학습자의 포트폴리오를 점검하는 취업 코치입니다. "
        "업로드 자료를 기준으로 분석하고, 백서는 프로젝트 공개 범위와 과정 설명 가이드로만 사용합니다. "
        "확인되지 않은 기업 정보나 성과를 지어내지 않습니다."
    )
    user = f"""
[백서 가이드]
{context}

[지원 직무]
{target_role or '미입력'}

[채용공고/기업 정보]
{job_text[:3500] if job_text else '미입력'}

[사용자 포트폴리오/프로젝트 자료]
{portfolio_text[:8000]}

[출력 형식]
### 포트폴리오 요약
### 강점
### 부족한 항목
### 직무별 보완 방향
### 면접에서 질문 나올 부분
### 공개·저작권 주의사항
"""
    response = ask_llm(system, user, temperature=0.2)
    if response:
        return response
    score_items = {
        "문제 정의": any(w in portfolio_text for w in ["문제", "배경", "목표", "니즈"]),
        "본인 역할": any(w in portfolio_text for w in ["역할", "담당", "기여", "구현", "분석"]),
        "기술 스택": any(w in portfolio_text.lower() for w in ["python", "sql", "ai", "ml", "rag", "cloud", "aws", "azure", "streamlit"]),
        "성과 표현": any(w in portfolio_text for w in ["성과", "%", "정확도", "개선", "결과", "효과"]),
        "회고/보완": any(w in portfolio_text for w in ["어려움", "한계", "개선", "보완", "배운"]),
    }
    missing = [k for k, v in score_items.items() if not v]
    present = [k for k, v in score_items.items() if v]
    return f"""### 포트폴리오 요약
업로드 자료에서 확인된 항목: {', '.join(present) if present else '제한적'}

### 강점
- 자료에 명시된 프로젝트 경험을 중심으로 정리 가능합니다.
- 지원 직무가 `{target_role or '미입력'}`로 설정되어 있어 직무 연관성을 맞춰 보완할 수 있습니다.

### 부족한 항목
{chr(10).join([f'- {item}' for item in missing]) if missing else '- 주요 항목이 대체로 포함되어 있습니다.'}

### 직무별 보완 방향
- 문제 정의 → 본인 역할 → 사용 기술 → 결과 → 배운 점 순서로 재정리하세요.
- 수치 성과가 없으면 처리 시간 단축, 사용자 편의, 재현 가능성 같은 정성 성과라도 분리해 적으세요.

### 면접에서 질문 나올 부분
- 이 프로젝트에서 본인이 직접 맡은 부분은 무엇인가요?
- 기술을 선택한 이유는 무엇인가요?
- 실패하거나 성능이 낮았던 부분을 어떻게 개선했나요?

### 공개·저작권 주의사항
- AIVLE 내부 교안, 실습자료, 과제 데이터, 산출물 공개 가능 범위는 백서와 공식 안내를 기준으로 확인하세요.
"""


def generate_interview_questions(portfolio_text: str, job_text: str, target_role: str) -> str:
    guide = career_whitepaper_guide("면접 포트폴리오 프로젝트 취업지원 잡페어", k=4)
    base = f"""
[백서 가이드 일부]
{guide or '백서 가이드 없음'}

[지원 직무]
{target_role or '미입력'}

[채용공고/기업 정보]
{job_text[:3500] if job_text else '미입력'}

[포트폴리오/프로젝트 자료]
{portfolio_text[:7000] if portfolio_text else '미입력'}
"""
    system = "업로드 자료를 바탕으로 면접 예상 질문, 꼬리 질문, 답변 방향을 만드는 취업 코치입니다. 없는 경험은 지어내지 않습니다."
    user = base + """
[출력 형식]
### 직무 공통 질문 5개
### 프로젝트 기반 질문 5개
### 꼬리 질문
### STAR 답변 틀
### 답변에서 피해야 할 표현
"""
    response = ask_llm(system, user, temperature=0.25)
    if response:
        return response
    return f"""### 직무 공통 질문 5개
1. {target_role or '지원 직무'}에 관심을 갖게 된 이유는 무엇인가요?
2. AIVLE 과정에서 가장 많이 성장한 역량은 무엇인가요?
3. 협업 과정에서 갈등이 있었을 때 어떻게 해결했나요?
4. 최근 학습한 기술을 실무 문제에 어떻게 적용할 수 있나요?
5. 본인의 강점과 보완점은 무엇인가요?

### 프로젝트 기반 질문 5개
1. 프로젝트의 문제 정의는 무엇이었나요?
2. 본인이 직접 맡은 역할은 무엇이었나요?
3. 기술 선택 기준은 무엇이었나요?
4. 결과를 어떻게 검증했나요?
5. 다시 한다면 무엇을 개선하겠나요?

### 꼬리 질문
- 왜 그 방식을 선택했나요?
- 다른 대안과 비교했나요?
- 수치로 설명할 수 있는 성과가 있나요?

### STAR 답변 틀
- Situation: 상황과 문제
- Task: 본인 역할
- Action: 직접 수행한 행동
- Result: 결과와 배운 점
"""


def improve_resume_text(raw_text: str, target_role: str) -> str:
    if not raw_text.strip():
        return "정리할 문장을 입력하거나 파일을 업로드해 주세요."
    guide = career_whitepaper_guide("포트폴리오 프로젝트 산출물 공개 제한 취업지원", k=3)
    system = "자기소개서, README, 포트폴리오 문장을 직무 중심으로 다듬는 코치입니다. 과장하거나 없는 성과를 만들지 않습니다."
    user = f"""
[지원 직무]
{target_role or '미입력'}

[백서 가이드 일부]
{guide or '백서 가이드 없음'}

[원문]
{raw_text[:6000]}

[출력 형식]
### 다듬은 문장
### 강조할 키워드
### 보완할 근거
### 면접 대비 메모
"""
    response = ask_llm(system, user, temperature=0.2)
    if response:
        return response
    return f"""### 다듬은 문장
{raw_text[:1000]}

### 강조할 키워드
- 지원 직무: {target_role or '미입력'}
- 문제 정의, 본인 역할, 사용 기술, 결과, 배운 점

### 보완할 근거
- 수치 성과 또는 검증 기준
- 팀 내 본인 기여
- 기술 선택 이유

### 면접 대비 메모
- 문장마다 실제 설명 가능한 근거를 준비하세요.
"""


def learning_coach_feedback(results: List[Dict[str, Any]], wrong_notes: List[Dict[str, Any]]) -> str:
    if not results and not wrong_notes:
        return "아직 진단 데이터가 없습니다. 예습·진단 메뉴에서 쪽지시험을 먼저 풀어 주세요."
    system = "쪽지시험 결과와 오답노트를 바탕으로 다음 학습 행동을 제안하는 코치입니다."
    user = f"""
[최근 진단 결과]
{json.dumps(results[-5:], ensure_ascii=False, indent=2)}

[최근 오답노트]
{json.dumps(wrong_notes[-10:], ensure_ascii=False, indent=2)}

[출력 형식]
### 현재 상태
### 취약 주제
### 다음 3일 학습 행동
### 추천 질문
"""
    response = ask_llm(system, user, temperature=0.2)
    if response:
        return response
    latest = results[-1] if results else {}
    level = latest.get("level", "미확인")
    weak = []
    for row in latest.get("stats", []):
        try:
            if float(row.get("정답률", 100)) < 70:
                weak.append(row.get("주제", "기타"))
        except Exception:
            pass
    return f"""### 현재 상태
- 최근 등급: **{level}**
- 누적 오답 수: **{len(wrong_notes)}개**

### 취약 주제
- {', '.join(weak) if weak else '최근 결과에서 뚜렷한 취약 주제가 확인되지 않았습니다.'}

### 다음 3일 학습 행동
1. 오답노트에서 같은 주제 문제를 다시 읽습니다.
2. 취약 주제를 학습 질의에 질문합니다.
3. 예습·진단 메뉴에서 같은 주제로 쪽지시험을 다시 풉니다.

### 추천 질문
- 이 주제의 핵심 개념을 쉬운 예시로 설명해줘
- 내가 틀린 개념과 연결되는 프로젝트 사례를 알려줘
"""


# ============================================================
# 대화 / 캘린더 / 체크리스트
# ============================================================


def load_conversations() -> Dict[str, Any]:
    data = read_json(CONVERSATION_PATH, {})
    if not isinstance(data, dict) or not data:
        cid = str(uuid.uuid4())
        data = {cid: {"title": "새 학습 대화", "created_at": datetime.now().isoformat(timespec="seconds"), "messages": []}}
        write_json(CONVERSATION_PATH, data)
    if st.session_state.get("conversation_id") not in data:
        st.session_state["conversation_id"] = next(iter(data.keys()))
    return data


def save_conversations(data: Dict[str, Any]) -> None:
    write_json(CONVERSATION_PATH, data)


def current_conversation_id() -> str:
    data = load_conversations()
    cid = st.session_state.get("conversation_id")
    if cid not in data:
        cid = next(iter(data.keys()))
        st.session_state["conversation_id"] = cid
    return cid


def create_conversation() -> str:
    data = load_conversations()
    cid = str(uuid.uuid4())
    data[cid] = {"title": "새 학습 대화", "created_at": datetime.now().isoformat(timespec="seconds"), "messages": []}
    save_conversations(data)
    st.session_state["conversation_id"] = cid
    return cid


def append_message(conv_id: str, role: str, content: str, sources: Optional[List[Dict[str, Any]]] = None) -> None:
    data = load_conversations()
    conv = data.setdefault(conv_id, {"title": "새 학습 대화", "created_at": datetime.now().isoformat(timespec="seconds"), "messages": []})
    conv["messages"].append({"role": role, "content": content, "sources": sources or [], "time": datetime.now().isoformat(timespec="seconds")})
    if role == "user" and conv.get("title") == "새 학습 대화":
        conv["title"] = content[:24] + ("..." if len(content) > 24 else "")
    save_conversations(data)


def read_calendar() -> List[Dict[str, Any]]:
    data = read_json(CALENDAR_PATH, [])
    if not isinstance(data, list):
        data = []
    default_events = [
        {"id": "default_1", "title": "AIVLE DAY 1차", "date": "2026-05-15", "kind": "시험", "note": "포트폴리오 강의, 코딩테스트"},
        {"id": "default_2", "title": "AIVLE DAY 2차", "date": "2026-06-26", "kind": "시험", "note": "면접 특강"},
        {"id": "default_3", "title": "Job-Fair", "date": "2026-08-27", "kind": "채용", "note": "취업지원 프로그램"},
    ]
    default_ids = {item["id"] for item in default_events}
    user_events = [item for item in data if item.get("id") not in default_ids]
    return default_events + user_events


def add_calendar_event(title: str, event_date: str, kind: str, note: str) -> None:
    data = read_json(CALENDAR_PATH, [])
    if not isinstance(data, list):
        data = []
    data.append({"id": str(uuid.uuid4()), "title": title, "date": event_date, "kind": kind, "note": note})
    write_json(CALENDAR_PATH, data)


def delete_calendar_event(event_id: str) -> None:
    data = read_json(CALENDAR_PATH, [])
    if not isinstance(data, list):
        return
    data = [item for item in data if item.get("id") != event_id]
    write_json(CALENDAR_PATH, data)


def load_checklist() -> Dict[str, Any]:
    data = read_json(CHECKLIST_PATH, {})
    return data if isinstance(data, dict) else {}


def save_checklist(data: Dict[str, Any]) -> None:
    write_json(CHECKLIST_PATH, data)


def save_career_report(title: str, report: str, report_type: str) -> None:
    data = read_json(CAREER_REPORT_PATH, [])
    if not isinstance(data, list):
        data = []
    data.append({"time": datetime.now().isoformat(timespec="seconds"), "type": report_type, "title": title, "report": report})
    write_json(CAREER_REPORT_PATH, data)


# ============================================================
# 렌더링 공통
# ============================================================


def hero(title: str, subtitle: str) -> None:
    html_block(f"""
    <div class='learning-hero'>
        <div class='hero-kicker'>학습 플랫폼</div>
        <h1>{safe_text(title)}</h1>
        <p>{safe_text(subtitle)}</p>
    </div>
    """)


def section(title: str, subtitle: str = "") -> None:
    st.markdown(f"<div class='section-title'>{safe_text(title)}</div>", unsafe_allow_html=True)
    if subtitle:
        st.markdown(f"<div class='section-sub'>{safe_text(subtitle)}</div>", unsafe_allow_html=True)


def small_help(text: str) -> None:
    st.caption(str(text))


def status_badge(text: str, kind: str = "info") -> None:
    class_name = {"success": "badge-success", "warning": "badge-warning", "muted": "badge-muted"}.get(kind, "badge-info")
    st.markdown(f"<span class='status-badge {class_name}'>{safe_text(text)}</span>", unsafe_allow_html=True)


def notice_card(title: str, body: str, badge: str = "안내", kind: str = "info") -> None:
    border_kind = "danger-zone" if kind == "danger" else "notice-card"
    html_block(f"""
    <div class='{border_kind} notice-card'>
        <span class='status-badge badge-info'>{safe_text(badge)}</span><br>
        <b>{safe_text(title)}</b>
        <p style='margin:.45rem 0 0;color:var(--muted);line-height:1.55;'>{safe_text(body)}</p>
    </div>
    """)


def empty_state(title: str, body: str) -> None:
    with st.container(border=True):
        st.markdown(f"**{title}**")
        st.caption(str(body))


def summary_cards(items: List[Tuple[str, str, str]]) -> None:
    """요약 카드를 Streamlit 네이티브 컨테이너로 표시합니다.
    HTML 렌더링 실패 시 코드가 그대로 보이는 문제를 방지합니다.
    """
    if not items:
        return
    cols = st.columns(min(len(items), 4))
    for idx, (label, value, hint) in enumerate(items):
        with cols[idx % len(cols)]:
            with st.container(border=True):
                st.caption(str(label))
                st.markdown(f"### {safe_text(value)}")
                st.caption(str(hint))


def get_record_counts() -> Dict[str, int]:
    """대시보드와 학습 현황에서 즉시 재사용할 누적 기록 수를 반환합니다."""
    conversations = load_conversations()
    results = read_json(RESULT_PATH, [])
    wrong_notes = read_json(WRONG_NOTE_PATH, [])
    events = read_json(CALENDAR_PATH, [])
    return {
        "conversations": len(conversations) if isinstance(conversations, dict) else 0,
        "results": len(results) if isinstance(results, list) else 0,
        "wrong_notes": len(wrong_notes) if isinstance(wrong_notes, list) else 0,
        "events": len(events) if isinstance(events, list) else 0,
    }


def mark_records_updated() -> None:
    """진단·오답 저장 직후 현재 실행 상태에서도 최신 기록임을 표시합니다."""
    st.session_state["records_updated_at"] = datetime.now().isoformat(timespec="seconds")
    st.session_state["record_counts"] = get_record_counts()


def horizontal_rate_chart(df: pd.DataFrame, label_col: str = "주제", value_col: str = "정답률") -> None:
    """긴 한국어 주제명이 세로로 깨지지 않도록 가로 막대 차트로 표시합니다."""
    if df is None or df.empty or label_col not in df.columns or value_col not in df.columns:
        empty_state("표시할 시각화 데이터가 없습니다", "쪽지시험을 채점하면 취약 주제 그래프가 표시됩니다.")
        return

    chart_df = df[[label_col, value_col]].copy()
    chart_df[label_col] = chart_df[label_col].astype(str).str.replace("\n", " ", regex=False)
    chart_df[value_col] = pd.to_numeric(chart_df[value_col], errors="coerce").fillna(0).clip(lower=0, upper=100)
    chart_df = chart_df.sort_values(value_col, ascending=True)

    if alt is None:
        st.dataframe(chart_df, hide_index=True, use_container_width=True)
        return

    chart = (
        alt.Chart(chart_df)
        .mark_bar(cornerRadiusEnd=6, height=18)
        .encode(
            x=alt.X(f"{value_col}:Q", scale=alt.Scale(domain=[0, 100]), title="정답률(%)"),
            y=alt.Y(f"{label_col}:N", sort=None, title=None, axis=alt.Axis(labelAngle=0, labelLimit=260)),
            tooltip=[alt.Tooltip(f"{label_col}:N", title="주제"), alt.Tooltip(f"{value_col}:Q", title="정답률", format=".1f")],
        )
        .properties(height=max(160, 42 * len(chart_df)))
    )
    st.altair_chart(chart, use_container_width=True)


def ensure_whitepaper_ready() -> bool:
    path = get_whitepaper_path()
    if not path.exists():
        notice_card("학습 자료가 없습니다", "왼쪽 사이드바에서 DOCX, PDF, TXT 형식의 백서나 커리큘럼 자료를 업로드하면 검색과 답변 기능을 사용할 수 있습니다.", badge="자료 필요")
        return False
    return True


def render_sources(hits: List[SearchHit]) -> None:
    if not hits:
        empty_state("표시할 백서 근거가 없습니다", "질문을 더 구체화하거나 학습 자료를 다시 업로드해 주세요.")
        return
    with st.expander("백서 근거 보기", expanded=False):
        for idx, hit in enumerate(hits, start=1):
            preview = safe_text(hit.text.replace("\n", " ")[:520])
            title = safe_text(hit.title)
            ellipsis = "..." if len(hit.text) > 520 else ""
            html_block(f"""
            <div class='source-box'>
                <b>근거 {idx} · {title}</b><br>
                <span style='color:#64748b;font-size:.86rem;'>유사도 {hit.score:.3f} · 청크 {hit.index}</span>
                <p>{preview}{ellipsis}</p>
            </div>
            """)


def render_link_table(links: List[Dict[str, str]]) -> None:
    if not links:
        return
    with st.expander("추천 학습 사이트", expanded=False):
        st.dataframe(pd.DataFrame(links), hide_index=True, use_container_width=True)


def render_header_metrics() -> None:
    index = load_index()
    chunk_count = len(index.get("chunks", [])) if index else 0
    counts = get_record_counts()
    meta = get_whitepaper_meta()
    display_name = meta.get("display_name") or get_whitepaper_path().name
    material_status = "반영 완료" if chunk_count > 0 else "자료 필요"
    status_hint = "백서 기반 답변 사용 가능" if chunk_count > 0 else "학습 자료를 업로드해 주세요"
    summary_cards([
        ("현재 학습 자료", display_name, "업로드한 백서·커리큘럼 기준"),
        ("학습 자료 상태", material_status, status_hint),
        ("저장 대화", f"{counts['conversations']:,}", "이어볼 수 있는 학습 대화"),
        ("진단 / 오답", f"{counts['results']:,} / {counts['wrong_notes']:,}", "누적 학습 점검 기록"),
    ])


def upload_signature(uploaded: Any) -> str:
    payload = uploaded.getvalue()
    digest = hashlib.sha256(payload).hexdigest()[:16]
    return f"{uploaded.name}:{uploaded.size}:{digest}"


def replace_current_whitepaper(uploaded: Any) -> Tuple[Path, bytes]:
    suffix = Path(uploaded.name).suffix.lower()
    if suffix not in WHITEPAPER_EXTENSIONS:
        raise ValueError("Unsupported whitepaper file type")
    WHITEPAPER_DIR.mkdir(exist_ok=True)
    target = WHITEPAPER_DIR / f"{CURRENT_WHITEPAPER_STEM}{suffix}"
    for existing in WHITEPAPER_DIR.glob(f"{CURRENT_WHITEPAPER_STEM}.*"):
        if existing != target and is_managed_whitepaper_path(existing):
            existing.unlink()
    payload = uploaded.getvalue()
    target.write_bytes(payload)
    return target, payload


def clear_whitepaper_index_cache() -> None:
    try:
        build_index.clear()
    except Exception:
        pass


def render_login_page() -> None:
    st.markdown("<br><br>", unsafe_allow_html=True)
    left, center, right = st.columns([1, 1.15, 1])
    with center:
        html_block("""
        <div class='info-panel' style='padding:28px;'>
            <div class='hero-kicker'>학습자 로그인</div>
            <h2 style='margin:4px 0 8px 0;color:#111827;'>AIVLE 학습도우미</h2>
            <p style='color:#64748b;'>제공받은 계정으로 로그인해 학습 질의, 예습, 진단, 복습, 취업 준비를 이용하세요.</p>
        </div>
        """)
        with st.form("login_form"):
            login_id = st.text_input("아이디")
            login_pw = st.text_input("비밀번호", type="password")
            submitted = st.form_submit_button("로그인", use_container_width=True)
        if submitted:
            if authenticate(login_id, login_pw):
                st.session_state["authenticated"] = True
                st.session_state["login_error"] = ""
                set_auth_query_token(login_id)
                st.rerun()
            else:
                st.session_state["login_error"] = "아이디 또는 비밀번호가 올바르지 않습니다."
        if st.session_state.get("login_error"):
            notice_card("로그인 정보 확인", st.session_state["login_error"], badge="로그인 실패", kind="danger")


def render_sidebar() -> str:
    nav_target = st.session_state.get("nav_target")
    if nav_target in MENU_OPTIONS:
        st.session_state["active_page"] = nav_target
        # widget key는 radio 생성 전에만 갱신해야 StreamlitAPIException이 발생하지 않습니다.
        st.session_state["sidebar_page"] = nav_target
    st.session_state["nav_target"] = None

    st.sidebar.markdown(textwrap.dedent("""
    <div class='sidebar-brand'>
        <div class='brand-mark'>A</div>
        <div>
            <div class='brand-title'>AIVLE 학습도우미</div>
            <div class='brand-sub'>학습자 모드</div>
        </div>
    </div>
    """).strip(), unsafe_allow_html=True)

    meta = get_whitepaper_meta()
    path = get_whitepaper_path()
    material_status = "적용됨" if path.exists() else "없음"
    badge_class = "badge-success" if path.exists() else "badge-warning"
    st.sidebar.markdown(
        f"<span class='status-badge {badge_class}'>학습 자료 {material_status}</span>",
        unsafe_allow_html=True,
    )
    st.sidebar.caption(f"현재 자료: {meta.get('display_name', path.name)}")

    current_page = st.session_state.get("active_page", "대시보드")
    if current_page not in MENU_OPTIONS:
        current_page = "대시보드"
    if st.session_state.get("sidebar_page") not in MENU_OPTIONS:
        st.session_state["sidebar_page"] = current_page
    page = st.sidebar.radio(
        "학습 메뉴",
        MENU_OPTIONS,
        key="sidebar_page",
        on_change=sync_sidebar_page,
    )
    if page not in MENU_OPTIONS:
        page = "대시보드"
    st.session_state["active_page"] = page
    st.sidebar.markdown(f"<div class='current-pill'>현재 화면 · {page}</div>", unsafe_allow_html=True)

    if st.sidebar.button("새 학습 대화", use_container_width=True):
        create_conversation()
        request_nav("학습 질의")

    conversations = load_conversations()
    conv_ids = list(conversations.keys())
    current_conv = current_conversation_id()
    selected = st.sidebar.selectbox(
        "학습 대화 목록",
        conv_ids,
        index=conv_ids.index(current_conv) if current_conv in conv_ids else 0,
        format_func=lambda cid: conversations.get(cid, {}).get("title", "대화"),
    )
    st.session_state["conversation_id"] = selected

    if st.sidebar.button("현재 학습 대화 삭제", use_container_width=True):
        data = load_conversations()
        if len(data) <= 1:
            data[selected]["messages"] = []
            data[selected]["title"] = "새 학습 대화"
        else:
            data.pop(selected, None)
            st.session_state["conversation_id"] = next(iter(data.keys()))
        save_conversations(data)
        st.rerun()

    st.sidebar.markdown("<div class='side-label'>학습 자료</div>", unsafe_allow_html=True)
    uploaded = st.sidebar.file_uploader(
        "백서 / 커리큘럼 업로드",
        type=["docx", "pdf", "txt"],
        key="whitepaper_uploader",
    )
    st.sidebar.caption("DOCX/PDF/TXT 형식의 백서, 커리큘럼, FAQ, 수업 안내 자료를 올리면 앱의 검색 기준으로 사용됩니다.")
    if uploaded is not None:
        sig = upload_signature(uploaded)
        if sig != st.session_state.get("last_whitepaper_upload_sig"):
            preview_text = parse_uploaded_file(uploaded, max_chars=1500)
            if not preview_text.strip():
                st.session_state["last_whitepaper_upload_sig"] = sig
                st.sidebar.warning("텍스트를 추출할 수 없는 파일입니다. DOCX, 텍스트가 포함된 PDF, TXT 자료를 다시 올려 주세요.")
            else:
                target, payload = replace_current_whitepaper(uploaded)
                save_whitepaper_meta(uploaded.name, target, len(payload))
                st.session_state["last_whitepaper_upload_sig"] = sig
                clear_whitepaper_index_cache()
                st.sidebar.success("학습 자료가 반영되었습니다.")
                st.rerun()

    st.sidebar.divider()
    if st.sidebar.button("로그아웃", use_container_width=True):
        st.session_state["authenticated"] = False
        st.session_state["login_error"] = ""
        clear_auth_query_token()
        st.rerun()

    return page


# ============================================================
# 페이지
# ============================================================


def page_dashboard() -> None:
    hero("AIVLE 학습도우미", "질문, 예습, 진단, 복습, 취업 준비, 일정을 한 흐름으로 연결한 학습자용 화면입니다.")
    render_header_metrics()

    section("오늘의 학습 체크리스트", "하루 학습 흐름을 짧게 확인합니다.")
    today = date.today().isoformat()
    checklist_data = load_checklist()
    saved = checklist_data.get(today, {}) if isinstance(checklist_data.get(today, {}), dict) else {}
    tasks = [
        "추천 질문 1개 확인",
        "예습 자료 1개 생성",
        "쪽지시험 1회 풀이",
        "오답노트 확인",
        "캘린더 일정 확인",
    ]
    cols = st.columns(len(tasks))
    current_checks = {}
    for idx, task in enumerate(tasks):
        current_checks[task] = cols[idx].checkbox(task, value=bool(saved.get(task, False)), key=f"check_{today}_{idx}")
    done_count = sum(1 for v in current_checks.values() if v)
    progress = done_count / len(tasks)
    st.progress(progress, text=f"오늘 진행률 {done_count}/{len(tasks)}")
    if st.button("체크리스트 저장", use_container_width=True):
        checklist_data[today] = current_checks
        save_checklist(checklist_data)
        st.success("오늘의 체크리스트가 저장되었습니다.")

    section("학습 루틴", "처음 사용하는 경우 아래 순서대로 진행하면 됩니다.")
    html_block("""
    <div class='routine-strip'>
        <div class='routine-item'><span>STEP 01</span><b>질문하기</b><p>추천 질문으로 막힌 개념을 빠르게 확인합니다.</p></div>
        <div class='routine-item'><span>STEP 02</span><b>예습하기</b><p>주차와 주제를 선택해 수업 전 핵심 내용을 정리합니다.</p></div>
        <div class='routine-item'><span>STEP 03</span><b>진단하기</b><p>쪽지시험으로 현재 이해도를 점수와 등급으로 확인합니다.</p></div>
        <div class='routine-item'><span>STEP 04</span><b>취업 연결</b><p>포트폴리오와 면접 질문으로 학습 경험을 정리합니다.</p></div>
    </div>
    """)

    section("바로 시작")
    c1, c2, c3 = st.columns(3)
    with c1:
        st.markdown("<div class='learn-card'><div class='badge'>질문</div><h3>백서 기반 질의응답</h3><p>추천 질문 또는 직접 질문으로 필요한 내용을 바로 확인합니다.</p></div>", unsafe_allow_html=True)
        if st.button("질문 화면으로 이동", use_container_width=True):
            st.session_state["pending_chat"] = "에이블스쿨 전체 학습 흐름을 요약해줘"
            request_nav("학습 질의")
    with c2:
        st.markdown("<div class='learn-card'><div class='badge'>진단</div><h3>예습·쪽지시험</h3><p>커리큘럼 주차별 예습 자료를 만들고 시험으로 수준을 확인합니다.</p></div>", unsafe_allow_html=True)
        if st.button("예습·진단으로 이동", use_container_width=True):
            request_nav("예습·진단")
    with c3:
        st.markdown("<div class='learn-card'><div class='badge'>취업</div><h3>포트폴리오·면접</h3><p>사용자 업로드 자료를 기반으로 취업 준비물을 정리합니다.</p></div>", unsafe_allow_html=True)
        if st.button("취업 준비로 이동", use_container_width=True):
            request_nav("취업 준비")

    section("최종 기능 구조", "겹치는 기능은 결과 화면 안으로 통합했습니다.")
    flow = pd.DataFrame(
        [
            ["학습 질의", "추천 질문 / 직접 질문 / 사이트 링크", "백서 근거로 빠르게 이해"],
            ["예습·진단", "예습 자료 / 쪽지시험 / 등급 판별", "수업 전후 이해도 확인"],
            ["복습·분석", "취약 주제 / 오답노트 / AI 코치", "약점 보완"],
            ["취업 준비", "포트폴리오 / 면접 / 채용공고 분석", "학습 경험을 취업 자료로 연결"],
            ["일정·커리큘럼", "캘린더 / 플래너 / 커리큘럼", "학습 계획 관리"],
        ],
        columns=["메뉴", "포함 기능", "목표"],
    )
    st.dataframe(flow, hide_index=True, use_container_width=True)


def page_chat() -> None:
    hero("학습 질의", "자주 묻는 질문을 버튼으로 선택하거나 직접 질문해 백서 기반 답변을 확인합니다.")
    status_badge("백서 근거 기반 답변", "success")
    status_badge("근거 부족 시 안내 카드 표시", "muted")
    if not ensure_whitepaper_ready():
        return

    section("추천 질문", "버튼을 누르면 바로 답변이 생성됩니다.")
    tabs = st.tabs(list(QUICK_QUESTION_GROUPS.keys()))
    for tab, (group, questions) in zip(tabs, QUICK_QUESTION_GROUPS.items()):
        with tab:
            cols = st.columns(3)
            for idx, question in enumerate(questions):
                if cols[idx % 3].button(question, key=f"quick_{group}_{idx}", use_container_width=True):
                    st.session_state["pending_chat"] = question
                    st.rerun()

    conv_id = current_conversation_id()
    conversations = load_conversations()
    messages = conversations.get(conv_id, {}).get("messages", [])
    st.divider()

    toolbar_left, toolbar_right = st.columns([3, 1])
    with toolbar_left:
        section("대화창", "대화는 고정된 박스 안에서만 쌓입니다. 페이지 전체가 길어지지 않도록 최근 대화부터 확인할 수 있습니다.")
    with toolbar_right:
        if st.button("대화 초기화", use_container_width=True, type="secondary"):
            conversations = load_conversations()
            conversations.setdefault(conv_id, {"title": "새 학습 대화", "messages": [], "created_at": datetime.now().isoformat(timespec="seconds")})
            conversations[conv_id]["messages"] = []
            conversations[conv_id]["title"] = "새 학습 대화"
            conversations[conv_id]["updated_at"] = datetime.now().isoformat(timespec="seconds")
            save_conversations(conversations)
            st.session_state.pop("pending_chat", None)
            st.rerun()

    if not messages:
        empty_state("아직 대화가 없습니다", "추천 질문을 누르거나 아래 입력창에 질문을 입력하면 백서 근거를 찾아 답변합니다.")

    st.caption("최근 30개 메시지만 대화창에 표시됩니다. 이전 내용은 대화 목록에 저장됩니다.")
    chat_box = st.container(height=520, border=True)
    with chat_box:
        for msg in messages[-30:]:
            role = msg.get("role", "assistant")
            with st.chat_message(role):
                st.markdown(msg.get("content", ""))
                sources = msg.get("sources") or []
                if sources:
                    try:
                        render_sources([SearchHit(**src) for src in sources])
                    except Exception:
                        notice_card("근거 표시 생략", "저장된 근거 형식이 맞지 않아 답변 본문만 표시합니다.", badge="근거")

    pending = st.session_state.pop("pending_chat", None)
    question = pending or st.chat_input("백서 기준으로 질문을 입력하세요")
    if question:
        append_message(conv_id, "user", question)
        with st.spinner("백서에서 근거를 찾고 답변을 생성하는 중입니다."):
            answer, hits, links = answer_from_whitepaper(question)
        append_message(conv_id, "assistant", answer, sources=[hit.__dict__ for hit in hits])
        st.rerun()


def page_prep_diagnosis() -> None:
    hero("예습·진단", "주차와 주제를 선택해 예습 자료를 만들고, 쪽지시험으로 이해도를 점검합니다.")
    if not ensure_whitepaper_ready():
        return

    left, right = st.columns([0.92, 1.08])
    with left:
        section("예습 조건")
        week_options = [f"{row['week']} · {row['topic']}" for row in DEFAULT_CURRICULUM]
        selected_week = st.selectbox("주차", week_options)
        topic = st.selectbox("주제", list(TOPIC_ALIASES.keys()))
        minutes = st.slider("예습 가능 시간", min_value=10, max_value=90, value=30, step=10)
        if st.button("예습 자료 생성", use_container_width=True):
            with st.spinner("예습 자료를 생성하는 중입니다."):
                material, hits = generate_prep_material(selected_week, topic, minutes)
            st.session_state["prep_material"] = material
            st.session_state["prep_topic"] = topic
            st.session_state["prep_hits"] = [hit.__dict__ for hit in hits]
            st.session_state["quiz"] = []
            st.session_state["last_quiz_result"] = None
            st.rerun()
    with right:
        section("예습 자료")
        if st.session_state.get("prep_material"):
            st.markdown(st.session_state["prep_material"])
            hits = [SearchHit(**item) for item in st.session_state.get("prep_hits", [])]
            render_sources(hits)
        else:
            empty_state("예습 자료가 아직 없습니다", "왼쪽에서 주차, 주제, 학습 가능 시간을 선택한 뒤 예습 자료를 생성하세요.")

    st.divider()
    section("쪽지시험")
    if st.button("현재 예습 자료로 쪽지시험 생성", use_container_width=True, disabled=not bool(st.session_state.get("prep_material"))):
        with st.spinner("쪽지시험을 생성하는 중입니다."):
            st.session_state["quiz"] = generate_quiz(st.session_state.get("prep_topic", "프로젝트 수행"), st.session_state.get("prep_material", ""), count=5)
            st.session_state["last_quiz_result"] = None
        st.rerun()

    quiz = st.session_state.get("quiz") or []
    if not quiz:
        return

    selected: Dict[int, int] = {}
    with st.form("quiz_form"):
        for idx, item in enumerate(quiz):
            choice = st.radio(
                f"Q{idx + 1}. {item['question']}",
                options=list(range(len(item["choices"]))),
                format_func=lambda i, item=item: item["choices"][i],
                key=f"quiz_{idx}_{hashlib.md5(item['question'].encode()).hexdigest()[:8]}",
            )
            selected[idx] = int(choice)
        submitted = st.form_submit_button("채점하고 수준 판별")

    if submitted:
        correct = 0
        wrong_items = []
        for idx, item in enumerate(quiz):
            is_correct = selected[idx] == item["answer_index"]
            correct += int(is_correct)
            if not is_correct:
                wrong_items.append({
                    "time": datetime.now().isoformat(timespec="seconds"),
                    "topic": item.get("topic", "기타"),
                    "question": item["question"],
                    "my_answer": item["choices"][selected[idx]],
                    "correct_answer": item["choices"][item["answer_index"]],
                    "explanation": item["explanation"],
                })
        score = round(correct / max(1, len(quiz)) * 100, 1)
        level = judge_level(score)
        stats = build_topic_stats(quiz, selected)
        weak_topics = stats.loc[stats["정답률"] < 70, "주제"].tolist() if not stats.empty else []
        result = {"time": datetime.now().isoformat(timespec="seconds"), "topic": st.session_state.get("prep_topic", "기타"), "score": score, "level": level, "stats": stats.to_dict(orient="records")}
        save_result(result)
        save_wrong_notes(wrong_items)
        mark_records_updated()
        st.session_state["last_quiz_result"] = {**result, "weak_topics": weak_topics, "wrong_items": wrong_items}
        st.rerun()

    result = st.session_state.get("last_quiz_result")
    if result:
        st.success(f"점수: {result['score']}점 / 학습 등급: {result['level']}")
        live_counts = get_record_counts()
        status_badge(f"진단 {live_counts['results']}건 · 오답 {live_counts['wrong_notes']}건 반영", "success")
        stats_df = pd.DataFrame(result.get("stats", []))
        if not stats_df.empty:
            horizontal_rate_chart(stats_df, "주제", "정답률")
        st.markdown(study_recommendation(result["level"], result.get("weak_topics", [])))
        if result.get("wrong_items"):
            section("방금 생성된 오답노트")
            for item in result["wrong_items"]:
                st.markdown(f"**{item['topic']}** · {item['question']}\n\n- 내 답: {item['my_answer']}\n- 정답: {item['correct_answer']}\n- 해설: {item['explanation']}")


def page_review_analysis() -> None:
    hero("복습·분석", "쪽지시험 결과를 누적해 점수 추이, 취약 주제, 오답, AI 학습 코치 피드백을 확인합니다.")
    results = read_json(RESULT_PATH, [])
    wrong_notes = read_json(WRONG_NOTE_PATH, [])

    section("학습 진단 요약")
    if not results:
        empty_state("진단 기록이 없습니다", "예습·진단 메뉴에서 쪽지시험을 풀면 점수와 취약 주제를 확인할 수 있습니다.")
    else:
        df = pd.DataFrame(results)
        c1, c2, c3 = st.columns(3)
        c1.metric("응시 횟수", len(df))
        c2.metric("평균 점수", f"{df['score'].mean():.1f}")
        c3.metric("최근 등급", str(df.iloc[-1].get("level", "-")))
        trend = df[["time", "score"]].copy()
        trend["time"] = pd.to_datetime(trend["time"], errors="coerce")
        if not trend.dropna().empty:
            st.line_chart(trend.dropna().set_index("time")["score"])

        topic_rows = []
        for row in results:
            for stat in row.get("stats", []):
                topic_rows.append(stat)
        if topic_rows:
            topic_df = pd.DataFrame(topic_rows)
            summary = topic_df.groupby("주제", as_index=False).agg({"정답률": "mean", "문항수": "sum"})
            summary["정답률"] = summary["정답률"].round(1)
            section("부족한 주제 시각화")
            horizontal_rate_chart(summary, "주제", "정답률")
            st.dataframe(summary, hide_index=True, use_container_width=True)

    st.divider()
    section("AI 학습 코치", "쪽지시험 결과와 오답노트를 종합해 다음 행동을 제안합니다.")
    if st.button("학습 코치 피드백 생성", use_container_width=True):
        with st.spinner("피드백을 생성하는 중입니다."):
            st.session_state["coach_feedback"] = learning_coach_feedback(results, wrong_notes)
    if st.session_state.get("coach_feedback"):
        st.markdown(st.session_state["coach_feedback"])

    st.divider()
    section("오답노트")
    if not wrong_notes:
        empty_state("오답노트가 비어 있습니다", "쪽지시험에서 틀린 문제가 생기면 자동으로 오답노트에 저장됩니다.")
        return
    topic_filter = st.selectbox("주제 필터", ["전체"] + sorted({item.get("topic", "기타") for item in wrong_notes}))
    filtered = wrong_notes if topic_filter == "전체" else [item for item in wrong_notes if item.get("topic") == topic_filter]
    for item in reversed(filtered[-30:]):
        with st.expander(f"{item.get('topic', '기타')} · {item.get('question', '')}"):
            st.markdown(f"- 내 답: {item.get('my_answer')}\n- 정답: {item.get('correct_answer')}\n- 해설: {item.get('explanation')}\n- 기록 시각: {item.get('time')}")
            render_link_table(link_recommendations(item.get("topic", "")))


def page_career() -> None:
    hero("취업 준비", "사용자 업로드 자료를 기반으로 포트폴리오, 면접, 채용공고를 한 화면에서 정리합니다.")
    tabs = st.tabs(["포트폴리오 분석", "면접 준비", "채용공고 분석", "문장 정리"])

    with tabs[0]:
        section("포트폴리오 파일 분석", "백서는 가이드로만 사용하고, 실제 분석은 사용자가 업로드한 자료를 기준으로 합니다.")
        target_role = st.text_input("지원 직무", value=st.session_state.get("career_target_role", "AI 개발자 / 데이터 분석가"), key="portfolio_role")
        portfolio_file = st.file_uploader("포트폴리오 / 프로젝트 정리표 업로드", type=["pdf", "docx", "txt", "xlsx", "csv"], key="portfolio_uploader")
        st.caption("PDF: 포트폴리오·이력서 / Excel·CSV: 프로젝트명, 기간, 역할, 사용 기술, 성과, 어려웠던 점이 들어간 정리표를 올리세요.")
        job_text_input = st.text_area("채용공고 또는 기업 정보", height=130, placeholder="지원하려는 공고 내용이 있으면 붙여넣으세요. 없으면 비워도 됩니다.", key="portfolio_job_text")
        st.caption("채용공고 PDF/DOCX/TXT 또는 직무 요구사항은 아래 채용공고 분석 탭에 올리거나 본문으로 붙여넣으면 포트폴리오 보완 방향이 더 정확해집니다.")
        portfolio_text = parse_uploaded_file(portfolio_file) if portfolio_file else ""
        job_text = job_text_input.strip()
        if portfolio_text:
            with st.expander("추출된 포트폴리오 내용 미리보기", expanded=False):
                st.text(portfolio_text[:2500])
        if st.button("포트폴리오 분석", use_container_width=True, disabled=not bool(portfolio_text.strip())):
            with st.spinner("포트폴리오를 분석하는 중입니다."):
                report = analyze_portfolio(portfolio_text, job_text, target_role)
            st.session_state["portfolio_report"] = report
            save_career_report("포트폴리오 분석", report, "portfolio")
        if st.session_state.get("portfolio_report"):
            st.markdown(st.session_state["portfolio_report"])

    with tabs[1]:
        section("면접 예상 질문", "프로젝트 자료와 공고 내용을 기반으로 질문·꼬리질문·답변 틀을 만듭니다.")
        interview_role = st.text_input("지원 직무", value=st.session_state.get("interview_role", "DX 컨설턴트 / 서비스 기획"), key="interview_role_input")
        project_text = st.text_area("프로젝트 요약", height=180, placeholder="프로젝트명, 문제 정의, 본인 역할, 사용 기술, 성과를 적으세요.", key="interview_project_text")
        interview_file = st.file_uploader("프로젝트 자료 업로드", type=["pdf", "docx", "txt", "xlsx", "csv"], key="interview_file")
        st.caption("발표자료 요약본, 포트폴리오 PDF, 프로젝트 정리 Excel을 올리면 프로젝트 기반 질문을 만들 수 있습니다.")
        interview_job = st.text_area("채용공고 / 기업 정보", height=130, placeholder="공고의 주요 업무, 자격요건, 우대사항을 붙여넣으세요.", key="interview_job_text")
        file_text = parse_uploaded_file(interview_file) if interview_file else ""
        combined_project = f"{project_text}\n{file_text}".strip()
        if st.button("면접 질문 생성", use_container_width=True, disabled=not bool(combined_project.strip() or interview_job.strip())):
            with st.spinner("면접 질문을 생성하는 중입니다."):
                report = generate_interview_questions(combined_project, interview_job, interview_role)
            st.session_state["interview_report"] = report
            save_career_report("면접 질문", report, "interview")
        if st.session_state.get("interview_report"):
            st.markdown(st.session_state["interview_report"])

    with tabs[2]:
        section("채용공고 / 공모전 분석", "공고를 붙여넣거나 파일로 올리면 마감일, 적합도, 다음 행동을 정리합니다.")
        interests = st.text_input("관심 분야", value="AI, DX, Cloud, 데이터 분석, 서비스 기획", key="job_interest")
        notice_text = st.text_area("공고 내용", height=220, placeholder="공고명, 지원 자격, 마감일, 제출물, URL 등을 붙여넣으세요.", key="job_notice_text")
        col_notice, col_contest = st.columns(2)
        with col_notice:
            notice_file = st.file_uploader("채용공고 파일 업로드", type=["pdf", "docx", "txt", "xlsx", "csv"], key="notice_uploader")
        with col_contest:
            contest_file = st.file_uploader("공모전 파일 업로드", type=["pdf", "docx", "txt", "xlsx", "csv"], key="contest_uploader")
        st.caption("채용공고, 공모전 안내문, 프로젝트 모집글을 PDF/DOCX/TXT/Excel로 올리거나 본문을 붙여넣으세요.")
        notice_file_text = parse_uploaded_file(notice_file) if notice_file else ""
        contest_file_text = parse_uploaded_file(contest_file) if contest_file else ""
        full_notice = f"{notice_text}\n{notice_file_text}\n{contest_file_text}".strip()
        if st.button("공고 분석", use_container_width=True, disabled=not bool(full_notice.strip())):
            st.session_state["notice_summary"] = summarize_notice(full_notice, interests)
        summary = st.session_state.get("notice_summary")
        if summary:
            st.dataframe(pd.DataFrame([summary]), hide_index=True, use_container_width=True)
            if summary.get("deadline") and summary.get("deadline") != "미확인":
                if st.button("마감일을 캘린더에 추가", use_container_width=True):
                    add_calendar_event(summary.get("title", "공고 마감"), summary["deadline"], summary.get("type", "공고"), summary.get("actions", ""))
                    st.success("캘린더에 추가되었습니다.")

    with tabs[3]:
        section("자기소개서 / README 문장 정리", "초안 문장을 직무 중심으로 정리합니다.")
        resume_role = st.text_input("지원 직무", value="AI 개발자", key="resume_role")
        raw_text = st.text_area("정리할 원문", height=230, placeholder="자기소개서 문장, GitHub README 초안, 포트폴리오 설명문을 붙여넣으세요.", key="resume_raw_text")
        resume_file = st.file_uploader("문장 파일 업로드", type=["pdf", "docx", "txt"], key="resume_file")
        st.caption("자기소개서 초안, README 초안, 포트폴리오 설명문을 PDF/DOCX/TXT로 올릴 수 있습니다.")
        resume_file_text = parse_uploaded_file(resume_file) if resume_file else ""
        combined = f"{raw_text}\n{resume_file_text}".strip()
        if st.button("문장 정리", use_container_width=True, disabled=not bool(combined.strip())):
            with st.spinner("문장을 정리하는 중입니다."):
                report = improve_resume_text(combined, resume_role)
            st.session_state["resume_report"] = report
            save_career_report("문장 정리", report, "resume")
        if st.session_state.get("resume_report"):
            st.markdown(st.session_state["resume_report"])


def build_learning_plan(topic: str, hours: int, days: int) -> pd.DataFrame:
    day_count = max(1, min(days, 14))
    hours = max(1, hours)
    actions = [
        "백서 근거 읽기",
        "핵심 용어 정리",
        "예습 자료 생성",
        "쪽지시험 풀이",
        "오답노트 복습",
        "프로젝트 적용 아이디어 작성",
        "학습 질의로 추가 질문",
    ]
    rows = []
    start = date.today()
    for i in range(day_count):
        rows.append({
            "날짜": (start + timedelta(days=i)).isoformat(),
            "주제": topic,
            "학습 시간": f"{max(1, round(hours / day_count, 1))}시간",
            "할 일": actions[i % len(actions)],
        })
    return pd.DataFrame(rows)



def calendar_kind_color(kind: str) -> str:
    palette = {
        "수업": "#2F6F5E",
        "시험": "#C9822B",
        "스터디": "#4F7C8A",
        "공모전": "#DFAE43",
        "채용": "#B94A48",
        "개인": "#6C7A89",
        "커리큘럼": "#2F6F5E",
        "학습": "#4F7C8A",
        "진단": "#C9822B",
        "취업": "#B94A48",
        "프로젝트": "#DFAE43",
        "미니프로젝트": "#2F6F5E",
        "빅프로젝트": "#3B6EA8",
        "오리엔테이션": "#6C7A89",
    }
    return palette.get(str(kind or ""), "#6C7A89")


def curriculum_event_color(kind: str, index: int) -> str:
    text = str(kind or "")
    if "미니프로젝트" in text:
        return ["#2F6F5E", "#3B6EA8"][index % 2]
    if "빅프로젝트" in text or "프로젝트" in text:
        return ["#3B6EA8", "#2F6F5E"][index % 2]
    if "수업" in text:
        return ["#D97706", "#7C3AED"][index % 2]
    if "오리엔테이션" in text:
        return "#6C7A89"
    if "취업" in text:
        return "#B94A48"
    return calendar_kind_color(text)


def curriculum_segment_kind(segment: str, row_kind: str) -> str:
    text = str(segment or "")
    base_kind = str(row_kind or "")
    if "빅프로젝트" in text:
        return "빅프로젝트"
    if "미니프로젝트" in text:
        return "미니프로젝트"
    if "입교식" in text or "오리엔테이션" in text:
        return "오리엔테이션"
    if "취업" in text or "수료식" in text:
        return "취업"
    if base_kind == "프로젝트":
        return "수업"
    return base_kind or "수업"


def curriculum_topic_segments(topic: str) -> List[Tuple[str, str]]:
    raw_parts = [part.strip() for part in re.split(r"\s*\+\s*", str(topic or "")) if part.strip()]
    if not raw_parts:
        return []
    segments: List[Tuple[str, str]] = []
    previous_subject = ""
    for part in raw_parts:
        title = part
        if part == "미니프로젝트" and previous_subject:
            title = f"{previous_subject} 미니프로젝트"
        kind = curriculum_segment_kind(title, "")
        if kind not in {"미니프로젝트", "빅프로젝트"}:
            previous_subject = re.sub(r"\s*(미니프로젝트|빅프로젝트)\s*", "", title).strip() or previous_subject
        segments.append((kind, title))
    return segments


def split_date_ranges(start: str, end: str, count: int) -> List[Tuple[str, str]]:
    if count <= 1:
        return [(start, end)]
    start_dt = date.fromisoformat(start)
    end_dt = date.fromisoformat(end)
    total_days = max(1, (end_dt - start_dt).days)
    ranges: List[Tuple[str, str]] = []
    cursor = start_dt
    for idx in range(count):
        remaining_segments = count - idx
        remaining_days = max(1, (end_dt - cursor).days)
        span = max(1, round(remaining_days / remaining_segments))
        next_cursor = end_dt if idx == count - 1 else min(end_dt, cursor + timedelta(days=span))
        ranges.append((cursor.isoformat(), next_cursor.isoformat()))
        cursor = next_cursor
    return ranges


def render_calendar_legend() -> None:
    html_block("""
    <div class='calendar-legend'>
        <span><i style='background:#D97706;'></i>수업</span>
        <span><i style='background:#7C3AED;'></i>수업 보조</span>
        <span><i style='background:#2F6F5E;'></i>미니프로젝트</span>
        <span><i style='background:#3B6EA8;'></i>프로젝트</span>
        <span><i style='background:#B94A48;'></i>취업·마감</span>
    </div>
    """)


def normalize_calendar_date(value: Any) -> Optional[str]:
    parsed = pd.to_datetime(value, errors="coerce")
    if pd.isna(parsed):
        return None
    return parsed.strftime("%Y-%m-%d")


def parse_curriculum_period(period: str, year: int = 2026) -> Tuple[Optional[str], Optional[str]]:
    matches = re.findall(r"(\d{1,2})\.(\d{1,2})", str(period or ""))
    if not matches:
        return None, None
    start_month, start_day = map(int, matches[0])
    end_month, end_day = map(int, matches[-1])
    start_date = date(year, start_month, start_day)
    end_date = date(year, end_month, end_day) + timedelta(days=1)
    return start_date.isoformat(), end_date.isoformat()


def curriculum_calendar_events(rows: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    events: List[Dict[str, Any]] = []
    for idx, row in enumerate(rows or []):
        start, end = parse_curriculum_period(str(row.get("period", "")))
        if not start:
            continue
        topic = str(row.get("topic", "커리큘럼"))
        week = str(row.get("week", ""))
        row_kind = str(row.get("kind", "커리큘럼"))
        segments = curriculum_topic_segments(topic) or [(curriculum_segment_kind(topic, row_kind), topic)]
        ranges = split_date_ranges(start, end, len(segments))
        for segment_idx, ((segment_kind, segment_title), (segment_start, segment_end)) in enumerate(zip(segments, ranges)):
            color = curriculum_event_color(segment_kind, idx + segment_idx)
            events.append({
                "id": f"curriculum_{idx}_{segment_idx}",
                "title": f"{week} · {segment_title}".strip(" ·"),
                "date": segment_start,
                "end": segment_end,
                "kind": segment_kind or row_kind or "커리큘럼",
                "note": str(row.get("note", "")),
                "color": color,
                "week": week,
                "period": str(row.get("period", "")),
            })
    return events


def learning_plan_calendar_events(plan: pd.DataFrame) -> List[Dict[str, Any]]:
    events: List[Dict[str, Any]] = []
    if not isinstance(plan, pd.DataFrame) or plan.empty:
        return events
    for idx, row in plan.iterrows():
        event_date = normalize_calendar_date(row.get("날짜"))
        if not event_date:
            continue
        topic = str(row.get("주제", "학습"))
        todo = str(row.get("할 일", ""))
        events.append({
            "id": f"plan_{idx}",
            "title": f"학습: {topic}",
            "date": event_date,
            "kind": "학습",
            "note": todo,
        })
    return events


def diagnosis_calendar_events(results: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    events: List[Dict[str, Any]] = []
    for idx, item in enumerate(results or []):
        event_date = normalize_calendar_date(item.get("time"))
        if not event_date:
            continue
        topic = str(item.get("topic", "진단"))
        score = item.get("score", "-")
        level = str(item.get("level", ""))
        events.append({
            "id": f"diagnosis_{idx}",
            "title": f"진단: {topic} {score}점",
            "date": event_date,
            "kind": "진단",
            "note": level,
        })
    return events


def career_report_calendar_events(reports: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    events: List[Dict[str, Any]] = []
    for idx, item in enumerate(reports or []):
        event_date = normalize_calendar_date(item.get("time"))
        if not event_date:
            continue
        report_type = str(item.get("type", "취업"))
        title = str(item.get("title", report_type))
        events.append({
            "id": f"career_{idx}",
            "title": f"취업: {title}",
            "date": event_date,
            "kind": "취업",
            "note": report_type,
        })
    return events


def to_calendar_component_events(events: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    component_events: List[Dict[str, Any]] = []
    for item in events or []:
        event_date = str(item.get("date", ""))[:10]
        if not re.match(r"^\d{4}-\d{2}-\d{2}$", event_date):
            continue
        kind = str(item.get("kind", "일정"))
        title = str(item.get("title", "일정"))
        color = str(item.get("color") or calendar_kind_color(kind))
        component_event = {
            "id": str(item.get("id", uuid.uuid4())),
            "title": f"[{kind}] {title}",
            "start": event_date,
            "allDay": True,
            "backgroundColor": color,
            "borderColor": color,
            "textColor": "#FFFFFF",
        }
        end_date = str(item.get("end", ""))[:10]
        if re.match(r"^\d{4}-\d{2}-\d{2}$", end_date):
            component_event["end"] = end_date
        component_events.append(component_event)
    return component_events


def render_fallback_month_calendar(events: List[Dict[str, Any]], year: int, month: int) -> None:
    event_map: Dict[str, List[Dict[str, Any]]] = {}
    for item in events or []:
        event_date = str(item.get("date", ""))[:10]
        event_map.setdefault(event_date, []).append(item)

    cal = py_calendar.Calendar(firstweekday=6)
    weeks = cal.monthdatescalendar(year, month)
    headers = ["일", "월", "화", "수", "목", "금", "토"]
    rows = ["<table class='calendar-fallback'><thead><tr>" + "".join(f"<th>{h}</th>" for h in headers) + "</tr></thead><tbody>"]
    for week in weeks:
        rows.append("<tr>")
        for day in week:
            day_key = day.isoformat()
            muted = " muted-day" if day.month != month else ""
            pills = []
            for event in event_map.get(day_key, [])[:3]:
                title = safe_text(event.get("title", "일정"))
                kind = safe_text(event.get("kind", "일정"))
                color = safe_text(event.get("color") or calendar_kind_color(str(event.get("kind", ""))))
                pills.append(f"<span class='calendar-event-pill' style='background:{color};border-color:{color};color:#fff;'>[{kind}] {title}</span>")
            if len(event_map.get(day_key, [])) > 3:
                pills.append(f"<span class='calendar-event-pill'>+{len(event_map[day_key]) - 3}개</span>")
            rows.append(f"<td class='{muted}'><div class='day-num'>{day.day}</div>{''.join(pills)}</td>")
        rows.append("</tr>")
    rows.append("</tbody></table>")
    html_block("".join(rows))


def render_month_calendar(events: List[Dict[str, Any]], year: int, month: int, key_prefix: str = "calendar_component") -> None:
    component_events = to_calendar_component_events(events)
    initial_date = f"{year:04d}-{month:02d}-01"
    render_fallback_month_calendar(events, year, month)
    if streamlit_calendar is not None:
        try:
            with st.expander("인터랙티브 캘린더", expanded=False):
                streamlit_calendar(
                    events=component_events,
                    options={
                        "initialView": "dayGridMonth",
                        "initialDate": initial_date,
                        "locale": "ko",
                        "height": 620,
                        "headerToolbar": {
                            "left": "prev,next today",
                            "center": "title",
                            "right": "dayGridMonth",
                        },
                        "buttonText": {"today": "오늘", "month": "월"},
                    },
                    key=f"{key_prefix}_{year}_{month}",
                )
        except Exception:
            pass


def available_calendar_months(events: List[Dict[str, Any]]) -> List[str]:
    months = {date.today().strftime("%Y-%m")}
    for item in events or []:
        event_date = str(item.get("date", ""))[:10]
        if re.match(r"^\d{4}-\d{2}-\d{2}$", event_date):
            months.add(event_date[:7])
    return sorted(months)

def page_plan_curriculum() -> None:
    hero("일정·커리큘럼", "커리큘럼, 학습 플래너, 캘린더를 한 화면에서 관리합니다.")
    tabs = st.tabs(["커리큘럼", "학습 플래너", "캘린더"])

    with tabs[0]:
        section("커리큘럼 달력")
        df = pd.DataFrame(DEFAULT_CURRICULUM)
        c1, c2 = st.columns([.8, 1.2])
        with c1:
            kind_filter = st.multiselect("구분", sorted(df["kind"].unique()), default=sorted(df["kind"].unique()))
        with c2:
            keyword = st.text_input("커리큘럼 검색", placeholder="예: 생성형 AI, 빅프로젝트, Cloud")
        filtered = df[df["kind"].isin(kind_filter)]
        if keyword:
            filtered = filtered[filtered.apply(lambda row: keyword.lower() in " ".join(map(str, row.values)).lower(), axis=1)]
        curriculum_events = curriculum_calendar_events(filtered.to_dict("records"))
        if curriculum_events:
            month_options = available_calendar_months(curriculum_events)
            default_month = date.today().strftime("%Y-%m")
            default_index = month_options.index(default_month) if default_month in month_options else 0
            selected_month = st.selectbox("표시할 월", month_options, index=default_index, key="curriculum_calendar_month")
            year, month = map(int, selected_month.split("-"))
            render_calendar_legend()
            render_month_calendar(curriculum_events, year, month, key_prefix="curriculum_calendar")
            event_view = pd.DataFrame(curriculum_events)
            st.dataframe(event_view[["week", "period", "kind", "title", "date", "end", "note"]], hide_index=True, use_container_width=True)
        else:
            empty_state("표시할 커리큘럼이 없습니다", "검색어 또는 구분 필터를 조정해 주세요.")
        st.divider()
        section("백서에서 추가 확인")
        query = st.text_input("백서 검색어", value=keyword or "커리큘럼")
        if st.button("백서 근거 검색", use_container_width=True):
            render_sources(search_whitepaper(query, k=8))

    with tabs[1]:
        section("학습 플래너", "목표와 가능 시간을 입력하면 짧은 학습 계획을 만듭니다.")
        topic = st.selectbox("학습 목표", ["분석형 AI", "생성형 AI", "Cloud", "서비스 개발/제안", "프로젝트 수행", "취업 준비"])
        col1, col2 = st.columns(2)
        hours = col1.slider("총 학습 가능 시간", 1, 30, 6)
        days = col2.slider("계획 기간", 1, 14, 5)
        if st.button("학습 계획 생성", use_container_width=True):
            st.session_state["learning_plan"] = build_learning_plan(topic, hours, days)
        if isinstance(st.session_state.get("learning_plan"), pd.DataFrame):
            plan = st.session_state["learning_plan"]
            plan_events = learning_plan_calendar_events(plan)
            if plan_events:
                month_options = available_calendar_months(plan_events)
                default_month = date.today().strftime("%Y-%m")
                default_index = month_options.index(default_month) if default_month in month_options else 0
                selected_month = st.selectbox("표시할 월", month_options, index=default_index, key="plan_calendar_month")
                year, month = map(int, selected_month.split("-"))
                render_month_calendar(plan_events, year, month, key_prefix="plan_calendar")
            st.dataframe(plan, hide_index=True, use_container_width=True)
            if st.button("계획을 캘린더에 추가", use_container_width=True):
                for _, row in plan.iterrows():
                    add_calendar_event(f"학습: {row['주제']}", row["날짜"], "스터디", row["할 일"])
                st.success("학습 계획이 캘린더에 추가되었습니다.")

    with tabs[2]:
        events = read_calendar()
        month_options = available_calendar_months(events)
        default_month = date.today().strftime("%Y-%m")
        default_index = month_options.index(default_month) if default_month in month_options else 0
        selected_month = st.selectbox("표시할 월", month_options, index=default_index, key="calendar_month_filter")
        year, month = map(int, selected_month.split("-"))

        section("월간 캘린더")
        render_calendar_legend()
        render_month_calendar(events, year, month, key_prefix="schedule_calendar")

        df = pd.DataFrame(events)
        if not df.empty:
            df["date"] = pd.to_datetime(df["date"], errors="coerce")
            display = df[df["date"].dt.strftime("%Y-%m") == selected_month].sort_values("date")
        else:
            display = pd.DataFrame()

        left, right = st.columns([0.78, 1.22])
        with left:
            section("일정 추가")
            with st.form("calendar_add"):
                title = st.text_input("일정명")
                event_date = st.date_input("날짜", value=date.today())
                kind = st.selectbox("구분", ["수업", "시험", "스터디", "공모전", "채용", "개인"])
                note = st.text_area("메모", height=90)
                ok = st.form_submit_button("일정 저장")
                if ok:
                    if title.strip():
                        add_calendar_event(title.strip(), event_date.isoformat(), kind, note)
                        st.success("일정이 저장되었습니다.")
                        st.rerun()
                    else:
                        notice_card("일정명 필요", "일정명을 입력한 뒤 저장해 주세요.", badge="입력 확인", kind="warning")

        with right:
            section("해당 월 일정")
            if display.empty:
                empty_state("해당 월 일정이 없습니다", "수업, 시험, 스터디, 공고 마감일을 추가해 학습 흐름을 관리하세요.")
            else:
                st.dataframe(display[["date", "kind", "title", "note"]], hide_index=True, use_container_width=True)
                deletable = display[~display["id"].astype(str).str.startswith("default_")]
                if not deletable.empty:
                    target = st.selectbox(
                        "삭제할 일정",
                        [""] + deletable["id"].astype(str).tolist(),
                        format_func=lambda eid: "선택 안 함" if not eid else deletable.loc[deletable["id"].astype(str) == str(eid), "title"].iloc[0],
                    )
                    if target and st.button("선택 일정 삭제"):
                        delete_calendar_event(target)
                        st.rerun()


def page_learning_status() -> None:
    hero("내 학습 현황", "저장된 대화, 진단 결과, 오답노트, 일정, 취업 준비 기록을 확인합니다.")
    render_header_metrics()

    section("현재 학습 자료")
    meta = get_whitepaper_meta()
    path = get_whitepaper_path()
    c1, c2 = st.columns(2)
    c1.markdown(f"**자료명**  \n{meta.get('display_name', path.name)}")
    c2.markdown(f"**반영 시각**  \n{meta.get('updated_at') or '기본 학습 자료'}")

    section("최근 진단 기록")
    results = read_json(RESULT_PATH, [])
    if results:
        st.dataframe(pd.DataFrame(results[-8:])[["time", "topic", "score", "level"]], hide_index=True, use_container_width=True)
    else:
        empty_state("최근 진단 기록이 없습니다", "예습·진단 메뉴에서 쪽지시험을 풀면 이곳에 기록됩니다.")

    events = read_calendar()
    reports = read_json(CAREER_REPORT_PATH, [])
    status_events = events + diagnosis_calendar_events(results) + career_report_calendar_events(reports)

    section("학습 현황 캘린더")
    if status_events:
        month_options = available_calendar_months(status_events)
        default_month = date.today().strftime("%Y-%m")
        default_index = month_options.index(default_month) if default_month in month_options else 0
        selected_month = st.selectbox("표시할 월", month_options, index=default_index, key="status_calendar_month")
        year, month = map(int, selected_month.split("-"))
        render_calendar_legend()
        render_month_calendar(status_events, year, month, key_prefix="status_calendar")

        df = pd.DataFrame(status_events)
        df["date"] = pd.to_datetime(df["date"], errors="coerce")
        display = df[df["date"].dt.strftime("%Y-%m") == selected_month].sort_values("date")
        if not display.empty:
            st.dataframe(display[["date", "kind", "title", "note"]], hide_index=True, use_container_width=True)
    else:
        empty_state("표시할 학습 기록이 없습니다", "진단, 일정, 취업 준비 기록이 생기면 달력에 함께 표시됩니다.")

    section("최근 취업 준비 기록")
    if reports:
        view = pd.DataFrame([{k: r.get(k) for k in ["time", "type", "title"]} for r in reports[-10:]])
        st.dataframe(view, hide_index=True, use_container_width=True)
    else:
        empty_state("취업 준비 기록이 없습니다", "취업 준비 메뉴에서 포트폴리오, 면접, 공고 분석을 실행하면 기록됩니다.")


# ============================================================
# 라우터
# ============================================================

PAGES = {
    "대시보드": page_dashboard,
    "학습 질의": page_chat,
    "예습·진단": page_prep_diagnosis,
    "복습·분석": page_review_analysis,
    "취업 준비": page_career,
    "일정·커리큘럼": page_plan_curriculum,
    "내 학습 현황": page_learning_status,
}


def main() -> None:
    init_state()
    if not st.session_state.get("authenticated"):
        render_login_page()
        return
    page = render_sidebar()
    PAGES.get(page, page_dashboard)()


if __name__ == "__main__":
    main()

